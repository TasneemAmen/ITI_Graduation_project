# ============================================================
# LTE KPI Degradation Analyzer - Generate Word Report
# ============================================================
# This file contains functions for generating Word reports.
# ============================================================

import numpy as np
import pandas as pd
from datetime import datetime

from KPI_Configuration import KPI_CONFIGS, DATE_COL, SITE_COL, CELL_COL

# Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


_ENH_EPS = 1e-9


def enhancement_potential_core(df, degraded_cell_ids, *, target_col, bad_direction,
                               metric_kind, weight_col=None, mode="fix",
                               recent_days=1, baseline_days=7, min_cells=2):
    """Reliable enhancement-potential engine (shared by report and dashboard).

    Answers: "How much would the network KPI improve if the degraded cells were
    restored to their own baseline (mode='fix'), keeping their users in the
    network?"  An optional mode='remove' models the cells (and their users/load)
    leaving the network instead.

    Returns:
        (value, meta) where value is a percentage (positive = improvement) or
        float('nan') when the result is NOT computable. meta['reason'] explains a
        NaN. The function NEVER returns 0.0 to mean "could not compute" - 0.0 only
        ever means a genuine zero enhancement.

    Args:
        target_col:   target KPI column name.
        bad_direction:"low" or "high".
        metric_kind:  'extensive' -> aggregate as a network SUM (volumes/counts);
                      'intensive'  -> load-weighted MEAN (rates/%/throughput).
        weight_col:   per-cell load weight (active users / attempts) for intensive
                      KPIs. None -> unweighted mean.
        mode:         'fix' (restore degraded cells to baseline; users stay) or
                      'remove' (drop degraded cells and their load).
        recent_days:  size of the recent window (1 = last day).
        baseline_days:size of the baseline window immediately before the recent one.
        min_cells:    minimum distinct cells required on the recent day.
    """
    meta = {"reason": None, "n_cells": 0, "n_degraded_matched": 0,
            "metric_kind": metric_kind, "mode": mode, "weight_col": weight_col}

    if df is None or getattr(df, "empty", True):
        meta["reason"] = "no data"
        return float("nan"), meta
    if not target_col:
        meta["reason"] = "no target column configured"
        return float("nan"), meta
    for c in (DATE_COL, target_col, SITE_COL, CELL_COL):
        if c not in df.columns:
            meta["reason"] = f"missing column: {c}"
            return float("nan"), meta

    degraded_cell_ids = degraded_cell_ids or set()
    use_w = bool(weight_col) and weight_col in df.columns and metric_kind == "intensive"
    cols = [DATE_COL, SITE_COL, CELL_COL, target_col] + ([weight_col] if use_w else [])
    d = df[cols].copy()
    d[DATE_COL] = pd.to_datetime(d[DATE_COL], errors="coerce")
    d[target_col] = pd.to_numeric(d[target_col], errors="coerce")
    d = d.dropna(subset=[DATE_COL, target_col])
    if use_w:
        d[weight_col] = pd.to_numeric(d[weight_col], errors="coerce")
    if d.empty:
        meta["reason"] = "no valid rows after cleaning"
        return float("nan"), meta

    last_day = d[DATE_COL].max()
    recent_start = last_day - pd.Timedelta(days=recent_days - 1)
    base_end = recent_start - pd.Timedelta(days=1)
    base_start = base_end - pd.Timedelta(days=baseline_days - 1)
    recent = d[d[DATE_COL] >= recent_start]
    base = d[(d[DATE_COL] >= base_start) & (d[DATE_COL] <= base_end)]

    # Reduce per (cell, day) first, THEN average over days -> one value per cell.
    # This removes any distortion from uneven sampling (e.g. hourly rows).
    val_how = "sum" if metric_kind == "extensive" else "mean"

    def per_cell(frame, col, how):
        if frame.empty:
            return pd.Series(dtype=float)
        return frame.groupby([CELL_COL, DATE_COL])[col].agg(how).groupby(level=0).mean()

    recent_val = per_cell(recent, target_col, val_how)
    base_val = per_cell(base, target_col, val_how)
    cells = recent_val.index
    if len(cells) < min_cells:
        meta["reason"] = f"too few cells on recent day ({len(cells)})"
        return float("nan"), meta
    meta["n_cells"] = int(len(cells))

    if metric_kind == "intensive" and use_w:
        w = per_cell(recent, weight_col, "mean").reindex(cells).where(lambda s: s > 0)
        if int(w.notna().sum()) == 0:
            meta["reason"] = "no valid weights (all zero/NaN)"
            return float("nan"), meta
        keep = w.notna()
        cells = cells[keep]
        recent_val = recent_val[keep]
        base_val = base_val.reindex(cells)
        w = w[keep]
    else:
        w = pd.Series(1.0, index=cells)
        meta["weight_col"] = None if metric_kind == "intensive" else weight_col

    cell_site = recent.drop_duplicates(CELL_COL).set_index(CELL_COL)[SITE_COL].reindex(cells)
    deg_mask = pd.Series([(cell_site.get(c), c) in degraded_cell_ids for c in cells], index=cells)
    meta["n_degraded_matched"] = int(deg_mask.sum())
    if int(deg_mask.sum()) == 0:
        meta["reason"] = "no degraded cells matched recent-day data"
        return float("nan"), meta

    def network(values):
        v = values.values
        ww = w.reindex(values.index).values
        if metric_kind == "extensive":
            return float(np.nansum(v))
        sw = np.nansum(ww)
        return float("nan") if sw < _ENH_EPS else float(np.nansum(v * ww) / sw)

    before = network(recent_val)

    if mode == "fix":
        counterfactual = recent_val.copy()
        bv = base_val.reindex(cells)
        restore = deg_mask & bv.notna()
        if int(restore.sum()) == 0:
            meta["reason"] = "degraded cells have no baseline to restore to"
            return float("nan"), meta
        idx = restore.index[restore]
        counterfactual.loc[idx] = bv.loc[idx]
        after = network(counterfactual)
    else:  # remove
        survivors = recent_val[~deg_mask]
        if survivors.empty:
            meta["reason"] = "all cells degraded (remove-mode undefined)"
            return float("nan"), meta
        after = network(survivors)

    if before is None or np.isnan(before) or abs(before) < _ENH_EPS:
        meta["reason"] = "network baseline ~ 0 (ratio undefined)"
        return float("nan"), meta

    if bad_direction == "high":
        enh = ((before - after) / before) * 100.0
    else:
        enh = ((after - before) / before) * 100.0
    return float(enh), meta


def format_enhancement_pct(value):
    """Render an enhancement value for display: 'N/A' for NaN, else 'x.xx%'."""
    try:
        if value is None or (isinstance(value, float) and np.isnan(value)):
            return "N/A"
        return f"{value:.2f}%"
    except Exception:
        return "N/A"


def calculate_enhancement_potential(original_df, degraded_cell_ids, selected_kpi, mode="fix"):
    """Public wrapper used by reports. Looks up KPI metadata from KPI_CONFIGS and
    delegates to enhancement_potential_core().

    Returns a float percentage, or float('nan') when not computable (callers should
    render NaN via format_enhancement_pct(), not as a number).
    """
    config = KPI_CONFIGS.get(selected_kpi, {})
    value, _meta = enhancement_potential_core(
        original_df, degraded_cell_ids,
        target_col=config.get("target_kpi"),
        bad_direction=config.get("bad_direction", "low"),
        metric_kind=config.get("metric_kind", "intensive"),
        weight_col=config.get("weight_kpi"),
        mode=mode,
    )
    return value


def get_top_cells_by_degradation(output_df, n=10):
    """
    Get top N cells with highest degradation.
    
    Args:
        output_df: DataFrame with degraded cells
        n: Number of top cells to return
        
    Returns:
        DataFrame of top cells
    """
    if output_df is None or output_df.empty:
        return pd.DataFrame()
    
    sorted_df = output_df.sort_values('kpi_degradation_ratio_%', ascending=False, key=abs)
    return sorted_df.head(n)


def generate_word_report(output_df, summary_df, analysis_mode, selected_kpi, baseline_mode, 
                          enable_significance_test, save_path, original_df=None, 
                          degraded_cell_ids=None, log_callback=None):
    """
    Generate a Word report with analysis results.
    
    Args:
        output_df: Output DataFrame with degraded cells
        summary_df: Summary DataFrame (for all KPIs mode)
        analysis_mode: "single" or "all"
        selected_kpi: Currently selected KPI name
        baseline_mode: Baseline mode setting
        enable_significance_test: Whether t-test was enabled
        save_path: Path to save the report
        original_df: Original input DataFrame
        degraded_cell_ids: Set of (Site, Cell) tuples for degraded cells
        log_callback: Optional logging callback
        
    Returns:
        True if successful, False otherwise
    """
    def log_msg(msg):
        if log_callback:
            log_callback(msg)
    
    if not DOCX_AVAILABLE:
        log_msg("ERROR: python-docx not installed. Run: pip install python-docx")
        return False
    
    if output_df is None or (output_df.empty and (summary_df is None or summary_df.empty)):
        log_msg("ERROR: No results to generate report")
        return False
    
    try:
        doc = Document()
        
        # Title
        doc.add_heading('RF Optimization Analysis Report', 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("Developed by: Musketeers_Team (ITI Graduation Project 2026)")
        doc.add_paragraph(f"Version: 2.0 Enhanced")
        
        # Analysis Summary
        doc.add_heading('Analysis Summary', level=1)
        
        if analysis_mode == "all":
            doc.add_paragraph(f"Analysis Mode: All KPIs Analysis")
            doc.add_paragraph(f"Baseline Mode: {baseline_mode}")
            doc.add_paragraph(f"Significance Test: {'Enabled' if enable_significance_test else 'Disabled'}")
            if summary_df is not None and not summary_df.empty:
                doc.add_paragraph(f"Total KPIs Analyzed: {len(summary_df)}")
                doc.add_paragraph(f"Total Degraded Cells: {int(summary_df['degraded_cells_count'].sum())}")
        else:
            doc.add_paragraph(f"Analysis Mode: Single KPI Analysis")
            doc.add_paragraph(f"Selected KPI: {selected_kpi}")
            doc.add_paragraph(f"Baseline Mode: {baseline_mode}")
            doc.add_paragraph(f"Degraded Cells: {len(output_df) if output_df is not None else 0}")
            if original_df is not None and degraded_cell_ids:
                enhancement = calculate_enhancement_potential(original_df, degraded_cell_ids, selected_kpi)
                doc.add_paragraph(f"Enhancement Potential: {format_enhancement_pct(enhancement)}")
        
        # ============================================================
        # MAIN KPI SUMMARY TABLE - All KPIs with key metrics
        # ============================================================
        if analysis_mode == "all" and summary_df is not None and not summary_df.empty:
            doc.add_heading('KPI Analysis Summary', level=1)
            
            # Columns from summary_df data
            data_cols = ['kpi_name', 'degraded_cells_count', 'max_degradation_%', 'threshold_%']
            available_data_cols = [c for c in data_cols if c in summary_df.columns]
            
            # Enhancement is always calculated separately (not from summary_df)
            always_include_enhancement = original_df is not None and degraded_cell_ids
            
            # Total columns = data columns + enhancement (if available)
            all_table_cols = available_data_cols + (['enhancement_potential_%'] if always_include_enhancement else [])
            
            if all_table_cols:
                # Calculate enhancement potential for each KPI
                enhancement_values = {}
                if always_include_enhancement:
                    for _, row in summary_df.iterrows():
                        kpi_name = row.get('kpi_name')
                        enhancement_values[kpi_name] = calculate_enhancement_potential(
                            original_df, degraded_cell_ids, kpi_name
                        )
                
                # Create table with header row
                table = doc.add_table(rows=1, cols=len(all_table_cols))
                table.style = 'Table Grid'
                
                # Format headers
                headers = table.rows[0].cells
                header_map = {
                    'kpi_name': 'KPI Name',
                    'degraded_cells_count': 'Degraded Cell Count',
                    'max_degradation_%': 'Max Degradation %',
                    'threshold_%': 'Threshold %',
                    'enhancement_potential_%': 'Enhancement Potential %'
                }
                for i, col in enumerate(all_table_cols):
                    headers[i].text = header_map.get(col, col.replace('_', ' ').title())
                
                # Add data rows
                for _, row in summary_df.iterrows():
                    cells = table.add_row().cells
                    for i, col in enumerate(all_table_cols):
                        if col == 'enhancement_potential_%':
                            # Use calculated enhancement potential
                            kpi_name = row.get('kpi_name')
                            if kpi_name in enhancement_values:
                                cells[i].text = format_enhancement_pct(enhancement_values[kpi_name])
                            else:
                                cells[i].text = "N/A"
                        else:
                            val = row.get(col, '')
                            if col == 'kpi_name':
                                cells[i].text = str(val)[:80]
                            elif col == 'degraded_cells_count':
                                cells[i].text = str(int(val)) if pd.notna(val) else "N/A"
                            elif col == 'threshold_%':
                                cells[i].text = f"{val:.2f}%" if pd.notna(val) else "N/A"
                            elif col.endswith('_%'):
                                cells[i].text = f"{abs(float(val)):.2f}%" if pd.notna(val) else "N/A"
                            else:
                                cells[i].text = "N/A" if pd.isna(val) else str(val)[:80]
        
        doc.save(save_path)
        log_msg(f"Report saved: {save_path}")
        return True
        
    except Exception as e:
        log_msg(f"Error generating report: {e}")
        return False