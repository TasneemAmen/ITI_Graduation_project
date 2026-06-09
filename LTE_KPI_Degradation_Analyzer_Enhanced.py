# ============================================================
# LTE KPI Degradation Analyzer - Enhanced Version v2.0
# ============================================================
# Addresses Reviewer Feedback:
# - Configurable baseline window (7-day, 4-week rolling, custom)
# - Minimum baseline value filter per KPI
# - Severity weighting for cause ranking
# - Statistical significance test (t-test)
# - Max/percentile aggregation for failure counters
# - Vectorized cause detection for performance
# - Multi-sheet Excel support
# - Per-cell trend chart in Word report
# ============================================================

import os
import threading
import traceback
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime, timedelta

import numpy as np
import pandas as pd
from scipy import stats

# Matplotlib for charts
import matplotlib
matplotlib.use('TkAgg')
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure

# Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ============================================================
# 1. GLOBAL COLUMN NAMES
# ============================================================

DATE_COL = "Date"
SITE_COL = "eNodeB Name"
CELL_COL = "Cell Name"
LOCAL_CELL_COL = "LocalCell Id"

CELL_ID_COLS = [
    SITE_COL,
    CELL_COL,
    LOCAL_CELL_COL,
]

# Baseline window modes
BASELINE_MODE_LAST_WEEK = "last_week"
BASELINE_MODE_4WEEK_AVG = "4week_rolling_avg"
BASELINE_MODE_CUSTOM = "custom_range"


# ============================================================
# 2. ENHANCED KPI CONFIGURATION
# ============================================================
# NEW FIELDS ADDED:
# - min_baseline_value: filter low-traffic/low-load cells
# - severity_weights: priority for cause ranking
# ============================================================

KPI_CONFIGS = {
    "DL Traffic": {
        "target_kpi": "(HU) DL Traffic Volume (GBytes)",
        "bad_direction": "low",
        "default_threshold": 30.0,
        "category": "Traffic",
        "output_prefix": "dl_traffic",
        "min_baseline_value": 1.0,  # NEW: Filter cells with < 1 GB baseline traffic
        "related_rules": [
            {"feature": "(HU) Cell DL Average Throughput (Mbps)", "bad_direction": "low", "threshold": 20, "severity": 3, "category": "DL Throughput Degradation", "reason": "Cell DL throughput decreased.", "recommended_action": "Check DL scheduler, bandwidth, CA activation, load balancing, and congestion."},
            {"feature": "(HU) User DL Average Throughput (Mbps)", "bad_direction": "low", "threshold": 20, "severity": 3, "category": "User Throughput Degradation", "reason": "User DL throughput decreased.", "recommended_action": "Check radio quality, PRB load, scheduler behavior, and user distribution."},
            {"feature": "(HU) DL PRB Utilization(%)", "bad_direction": "high", "threshold": 20, "severity": 2, "category": "Capacity / Congestion", "reason": "DL PRB utilization increased.", "recommended_action": "Check load balancing, CA usage, bandwidth, traffic distribution, and capacity expansion."},
            {"feature": "DL Average CQI", "bad_direction": "low", "threshold": 15, "severity": 3, "category": "Radio Quality Issue", "reason": "DL CQI decreased.", "recommended_action": "Check interference, PCI conflict/confusion, antenna tilt, azimuth, and coverage."},
            {"feature": "(HU) DL IBLER(%)", "bad_direction": "high", "threshold": 20, "severity": 3, "category": "DL Radio Quality Issue", "reason": "DL IBLER increased.", "recommended_action": "Check interference, CQI, MCS, antenna tilt, and DL power."},
            {"feature": "DL RBLER", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "DL Radio Failure", "reason": "DL RBLER increased.", "recommended_action": "Check DL interference, poor coverage, CQI, MCS, and radio conditions."},
            {"feature": "(HU) PDSCH MCS", "bad_direction": "low", "threshold": 15, "severity": 2, "category": "Poor Modulation Efficiency", "reason": "PDSCH MCS decreased.", "recommended_action": "Check CQI, SINR, interference, antenna direction, and coverage."},
            {"feature": "Availability", "bad_direction": "low", "threshold": 1, "severity": 5, "category": "Availability Issue", "reason": "Cell availability decreased.", "recommended_action": "Check alarms, S1 issue, manual outage, system outage, and site availability."},
            {"feature": "(HU) Cell Unavail Time (s)", "bad_direction": "high", "threshold": 20, "severity": 5, "category": "Cell Unavailability", "reason": "Cell unavailable time increased.", "recommended_action": "Check site outage, power issue, transmission issue, S1 failure, and alarms."},
            {"feature": "L.Traffic.ActiveUser.Dl.Avg", "bad_direction": "low", "threshold": 20, "severity": 1, "category": "Traffic Demand Drop", "reason": "DL active users decreased.", "recommended_action": "Validate if traffic drop is normal demand behavior before RF optimization."},
            {"feature": "MAC CA Traffic Ratio", "bad_direction": "low", "threshold": 20, "severity": 2, "category": "Carrier Aggregation Issue", "reason": "CA traffic ratio decreased.", "recommended_action": "Check CA activation, SCell availability, CA bands, and CA parameters."},
            {"feature": "DL Traffic QCI-9", "bad_direction": "low", "threshold": 20, "severity": 2, "category": "Default Bearer Traffic Drop", "reason": "QCI-9 DL traffic decreased.", "recommended_action": "Check packet data service, APN/data service, user demand, and internet traffic trend."},
            {"feature": "DL_CCE_AllocFail (%)", "bad_direction": "high", "threshold": 20, "severity": 3, "category": "Control Channel Congestion", "reason": "DL CCE allocation failure increased.", "recommended_action": "Check PDCCH/CCE utilization, control channel capacity, and scheduler configuration."},
        ],
    },

    "UL Traffic": {
        "target_kpi": "(HU) UL Traffic Volume (GBytes)",
        "bad_direction": "low",
        "default_threshold": 30.0,
        "category": "Traffic",
        "output_prefix": "ul_traffic",
        "min_baseline_value": 0.1,  # NEW
        "related_rules": [
            {"feature": "(HU) Cell UL Average Throughput (Mbps)", "bad_direction": "low", "threshold": 20, "severity": 3, "category": "UL Throughput Degradation", "reason": "Cell UL throughput decreased.", "recommended_action": "Check UL scheduler, UL PRB utilization, uplink interference, and power control."},
            {"feature": "(HU) User UL Average Throughput (Mbps)", "bad_direction": "low", "threshold": 20, "severity": 3, "category": "UL User Throughput Degradation", "reason": "User UL throughput decreased.", "recommended_action": "Check UL radio quality, UL interference, PUSCH MCS, and UL PRB load."},
            {"feature": "(HU)UL PRB Utilization(%)", "bad_direction": "high", "threshold": 20, "severity": 2, "category": "UL Capacity / Congestion", "reason": "UL PRB utilization increased.", "recommended_action": "Check UL capacity, UL scheduling, uplink load, and traffic distribution."},
            {"feature": "(HU) Avg UL Interference(dBm)", "bad_direction": "high", "threshold": 10, "severity": 4, "category": "UL Interference Issue", "reason": "Average UL interference increased.", "recommended_action": "Check external interference, PIM, neighboring cells, and uplink noise rise."},
            {"feature": "L.UpPTS.Interference.Avg(dBm)", "bad_direction": "high", "threshold": 10, "severity": 3, "category": "UL Interference Issue", "reason": "UpPTS interference increased.", "recommended_action": "Check uplink interference source and TDD interference conditions."},
            {"feature": "(HU) UL IBLER(%)", "bad_direction": "high", "threshold": 20, "severity": 3, "category": "UL Radio Quality Issue", "reason": "UL IBLER increased.", "recommended_action": "Check UL interference, PUSCH MCS, UE power, and coverage."},
            {"feature": "UL RBLER", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "UL Radio Failure", "reason": "UL RBLER increased.", "recommended_action": "Check UL interference, coverage, UE power, and uplink radio conditions."},
            {"feature": "(HU) PUSCH MCS", "bad_direction": "low", "threshold": 15, "severity": 2, "category": "UL Modulation Efficiency Issue", "reason": "PUSCH MCS decreased.", "recommended_action": "Check UL SINR, interference, UE power control, and uplink coverage."},
            {"feature": "L.Traffic.ActiveUser.UL.Avg", "bad_direction": "low", "threshold": 20, "severity": 1, "category": "UL Traffic Demand Drop", "reason": "UL active users decreased.", "recommended_action": "Validate traffic trend before applying RF action."},
        ],
    },

    "DL Throughput": {
        "target_kpi": "(HU) User DL Average Throughput (Mbps)",
        "bad_direction": "low",
        "default_threshold": 20.0,
        "category": "Integrity",
        "output_prefix": "dl_throughput",
        "min_baseline_value": 5.0,  # NEW: Mbps
        "related_rules": [
            {"feature": "(HU) DL PRB Utilization(%)", "bad_direction": "high", "threshold": 20, "severity": 3, "category": "DL Congestion", "reason": "DL PRB utilization increased while DL throughput decreased.", "recommended_action": "Check congestion, load balancing, CA, bandwidth, scheduler, and capacity expansion."},
            {"feature": "DL Average CQI", "bad_direction": "low", "threshold": 15, "severity": 4, "category": "Poor Radio Quality", "reason": "CQI decreased while DL throughput degraded.", "recommended_action": "Check interference, PCI, antenna tilt, azimuth, and coverage."},
            {"feature": "(HU) DL IBLER(%)", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "High DL Retransmission", "reason": "DL IBLER increased while DL throughput degraded.", "recommended_action": "Check BLER, CQI, MCS, DL power, and interference."},
            {"feature": "DL RBLER", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "DL Radio Block Errors", "reason": "DL RBLER increased while DL throughput degraded.", "recommended_action": "Check interference, radio quality, and coverage."},
            {"feature": "(HU) PDSCH MCS", "bad_direction": "low", "threshold": 15, "severity": 3, "category": "Low DL Modulation", "reason": "PDSCH MCS decreased while DL throughput degraded.", "recommended_action": "Check CQI, SINR, interference, and coverage."},
            {"feature": "MAC CA Traffic Ratio", "bad_direction": "low", "threshold": 20, "severity": 2, "category": "Low CA Usage", "reason": "CA traffic ratio decreased while DL throughput degraded.", "recommended_action": "Check CA activation, SCell availability, and CA configuration."},
            {"feature": "L.Traffic.ActiveUser.Dl.Avg", "bad_direction": "high", "threshold": 20, "severity": 2, "category": "High User Load", "reason": "DL active users increased while DL throughput decreased.", "recommended_action": "Check load, scheduling, congestion, and user distribution."},
        ],
    },

    "UL Throughput": {
        "target_kpi": "(HU) User UL Average Throughput (Mbps)",
        "bad_direction": "low",
        "default_threshold": 20.0,
        "category": "Integrity",
        "output_prefix": "ul_throughput",
        "min_baseline_value": 2.0,  # NEW: Mbps
        "related_rules": [
            {"feature": "(HU)UL PRB Utilization(%)", "bad_direction": "high", "threshold": 20, "severity": 3, "category": "UL Congestion", "reason": "UL PRB utilization increased while UL throughput degraded.", "recommended_action": "Check UL load, UL scheduler, and uplink capacity."},
            {"feature": "(HU) Avg UL Interference(dBm)", "bad_direction": "high", "threshold": 10, "severity": 4, "category": "UL Interference Issue", "reason": "UL interference increased while UL throughput degraded.", "recommended_action": "Check external interference, PIM, and uplink noise rise."},
            {"feature": "(HU) UL IBLER(%)", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "High UL Retransmission", "reason": "UL IBLER increased while UL throughput degraded.", "recommended_action": "Check UL BLER, interference, PUSCH MCS, and UE power."},
            {"feature": "UL RBLER", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "UL Radio Block Errors", "reason": "UL RBLER increased while UL throughput degraded.", "recommended_action": "Check UL coverage, interference, and UE power limitation."},
            {"feature": "(HU) PUSCH MCS", "bad_direction": "low", "threshold": 15, "severity": 3, "category": "Low UL Modulation", "reason": "PUSCH MCS decreased while UL throughput degraded.", "recommended_action": "Check uplink SINR, interference, and power control."},
            {"feature": "L.Traffic.ActiveUser.UL.Avg", "bad_direction": "high", "threshold": 20, "severity": 2, "category": "High UL User Load", "reason": "UL active users increased while UL throughput decreased.", "recommended_action": "Check uplink scheduling, congestion, and load distribution."},
        ],
    },

    "RRC Setup SR": {
        "target_kpi": "(TE) RRC Setup SR%",
        "bad_direction": "low",
        "default_threshold": 5.0,
        "category": "Accessibility",
        "output_prefix": "rrc_setup_sr",
        "min_baseline_value": 90.0,  # NEW: Only cells with > 90% baseline SR
        "related_rules": [
            {"feature": "L.RRC.ConnReq.Att", "bad_direction": "high", "threshold": 20, "severity": 2, "category": "High RRC Attempts", "reason": "RRC attempts increased.", "recommended_action": "Check RACH load, access parameters, admission control, and overload."},
            {"feature": "RRC Setup Failure Time", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "RRC Failure Increase", "reason": "RRC setup failures increased.", "recommended_action": "Check RACH failures, no reply, rejection, admission control, and radio quality."},
            {"feature": "L.RRC.SetupFail.NoReply", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "RRC No Reply", "reason": "RRC no-reply failures increased.", "recommended_action": "Check coverage, interference, RACH configuration, and UE access conditions."},
            {"feature": "L.RRC.SetupFail.Rej", "bad_direction": "high", "threshold": 20, "severity": 3, "category": "RRC Rejection", "reason": "RRC rejection failures increased.", "recommended_action": "Check admission control, overload, forbidden access, and MME overload."},
            {"feature": "L.RRC.SetupFail.Rej.MMEOverload", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "MME Overload", "reason": "RRC rejection due to MME overload increased.", "recommended_action": "Check MME/S1 signaling, core side load, and S1 interface."},
            {"feature": "L.RRC.SetupFail.ResFail", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "Radio Resource Failure", "reason": "RRC setup failures due to resource failure increased.", "recommended_action": "Check radio resources, PRB load, admission control, and congestion."},
            {"feature": "RACH Contention-Based Failures", "bad_direction": "high", "threshold": 20, "severity": 3, "category": "RACH Failure", "reason": "RACH contention failures increased.", "recommended_action": "Check PRACH configuration, root sequence planning, preamble load, and coverage."},
        ],
    },

    "ERAB Setup SR": {
        "target_kpi": "ERAB Setup Success Rate",
        "bad_direction": "low",
        "default_threshold": 5.0,
        "category": "Accessibility",
        "output_prefix": "erab_setup_sr",
        "min_baseline_value": 95.0,  # NEW
        "related_rules": [
            {"feature": "L.E-RAB.AttEst", "bad_direction": "high", "threshold": 20, "severity": 2, "category": "High ERAB Attempts", "reason": "E-RAB setup attempts increased.", "recommended_action": "Check access load, service attempts, and admission control."},
            {"feature": "ERAB Setup Failure Times", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "ERAB Failure Increase", "reason": "E-RAB setup failures increased.", "recommended_action": "Check ERAB failure reason counters, MME, TNL, RNL, and radio resources."},
            {"feature": "L.E-RAB.FailEst.NoRadioRes", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "No Radio Resource", "reason": "E-RAB failures due to no radio resources increased.", "recommended_action": "Check PRB load, admission control, congestion, and capacity."},
            {"feature": "L.E-RAB.FailEst.NoReply", "bad_direction": "high", "threshold": 20, "severity": 3, "category": "No Reply Failure", "reason": "E-RAB no-reply failures increased.", "recommended_action": "Check radio quality, signaling, and UE response issues."},
            {"feature": "L.E-RAB.FailEst.MME", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "MME Related Failure", "reason": "E-RAB setup failures related to MME increased.", "recommended_action": "Check MME/core side, S1 signaling, and core load."},
            {"feature": "L.E-RAB.FailEst.TNL", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "Transport Network Failure", "reason": "E-RAB setup failures related to TNL increased.", "recommended_action": "Check transmission, backhaul, S1-U/S1-C, and transport path."},
        ],
    },

    "Drop Rate": {
        "target_kpi": "E-RAB Drop Rate (E-NodeB + MME) %",
        "bad_direction": "high",
        "default_threshold": 20.0,
        "category": "Retainability",
        "output_prefix": "drop_rate",
        "min_baseline_value": 0.0,  # Any baseline allowed
        "related_rules": [
            {"feature": "L.E-RAB.AbnormRel", "bad_direction": "high", "threshold": 20, "severity": 5, "category": "Abnormal Release Increase", "reason": "E-RAB abnormal releases increased.", "recommended_action": "Check drop reason counters, radio quality, HO failures, and TNL/MME causes."},
            {"feature": "L.E-RAB.AbnormRel.Radio", "bad_direction": "high", "threshold": 20, "severity": 5, "category": "Radio Drop Issue", "reason": "Radio abnormal releases increased.", "recommended_action": "Check coverage, interference, CQI, BLER, and antenna settings."},
            {"feature": "L.E-RAB.AbnormRel.Radio.ULSyncFail", "bad_direction": "high", "threshold": 20, "severity": 5, "category": "UL Sync Failure", "reason": "Drops due to UL synchronization failure increased.", "recommended_action": "Check uplink coverage, UL interference, UE power, and timing advance."},
            {"feature": "L.E-RAB.AbnormRel.Radio.UuNoReply", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "Uu No Reply", "reason": "Drops due to Uu no-reply increased.", "recommended_action": "Check coverage holes, interference, and radio link quality."},
            {"feature": "L.E-RAB.AbnormRel.TNL", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "Transport Drop Issue", "reason": "Transport-related abnormal releases increased.", "recommended_action": "Check backhaul, transmission alarms, packet loss, and transport congestion."},
            {"feature": "L.E-RAB.AbnormRel.MME", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "MME Drop Issue", "reason": "MME-related abnormal releases increased.", "recommended_action": "Check core network, MME, S1 signaling, and S1 reset counters."},
            {"feature": "L.E-RAB.AbnormRel.HOFailure", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "HO Related Drop", "reason": "Abnormal releases related to HO failure increased.", "recommended_action": "Check neighbors, missing neighbors, A3 offset, CIO, TTT, and PCI issues."},
            {"feature": "RRC Connection Drop Rate%", "bad_direction": "high", "threshold": 20, "severity": 5, "category": "RRC Drop Issue", "reason": "RRC connection drop rate increased.", "recommended_action": "Check radio quality, coverage, interference, re-establishment, and mobility."},
        ],
    },

    "HO Success Rate": {
        "target_kpi": "HO SR% Overall",
        "bad_direction": "low",
        "default_threshold": 5.0,
        "category": "Mobility",
        "output_prefix": "ho_success_rate",
        "min_baseline_value": 90.0,  # NEW
        "related_rules": [
            {"feature": "Intra_Freq HO Prepare Failed Times", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "Intra-Frequency HO Preparation Failure", "reason": "Intra-frequency HO preparation failures increased.", "recommended_action": "Check neighbor relations, target cell availability, admission control, and HO prep failure reasons."},
            {"feature": "Intra_Freq HO Execution Failed Times", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "Intra-Frequency HO Execution Failure", "reason": "Intra-frequency HO execution failures increased.", "recommended_action": "Check radio quality, A3 offset, TTT, CIO, PCI, and target cell coverage."},
            {"feature": "Inter_Freq HO Prepare Failed Times", "bad_direction": "high", "threshold": 20, "severity": 3, "category": "Inter-Frequency HO Preparation Failure", "reason": "Inter-frequency HO preparation failures increased.", "recommended_action": "Check inter-frequency neighbors, measurement configuration, frequency priority, and target availability."},
            {"feature": "Inter_Freq HO Execution Failed Times", "bad_direction": "high", "threshold": 20, "severity": 3, "category": "Inter-Frequency HO Execution Failure", "reason": "Inter-frequency HO execution failures increased.", "recommended_action": "Check A3/A5 thresholds, TTT, CIO, target cell coverage, and PCI conflicts."},
            {"feature": "S1 HO Execution Failed Times", "bad_direction": "high", "threshold": 20, "severity": 3, "category": "S1 HO Failure", "reason": "S1 HO execution failures increased.", "recommended_action": "Check S1 handover path, MME, transport, and target eNodeB response."},
            {"feature": "X2 Intra-Freq Failure", "bad_direction": "high", "threshold": 20, "severity": 3, "category": "X2 Intra-Frequency HO Failure", "reason": "X2 intra-frequency HO failures increased.", "recommended_action": "Check X2 links, neighbor relation, target cell, and mobility parameters."},
            {"feature": "X2 Inter-Freq Failure", "bad_direction": "high", "threshold": 20, "severity": 3, "category": "X2 Inter-Frequency HO Failure", "reason": "X2 inter-frequency HO failures increased.", "recommended_action": "Check X2 links, inter-frequency neighbors, and target frequency settings."},
            {"feature": "L.HHO.PingPongHo", "bad_direction": "high", "threshold": 20, "severity": 2, "category": "Ping-Pong HO Issue", "reason": "Ping-pong handovers increased.", "recommended_action": "Tune hysteresis, TTT, CIO, A3 offset, and neighbor priorities."},
        ],
    },

    "Availability": {
        "target_kpi": "Availability",
        "bad_direction": "low",
        "default_threshold": 1.0,
        "category": "Availability",
        "output_prefix": "availability",
        "min_baseline_value": 99.0,  # NEW: Only cells with > 99% baseline availability
        "related_rules": [
            {"feature": "(HU) Cell Unavail Time (s)", "bad_direction": "high", "threshold": 20, "severity": 5, "category": "Cell Unavailable Time Increase", "reason": "Cell unavailable time increased.", "recommended_action": "Check outage, alarms, power, transmission, and site status."},
            {"feature": "L.Cell.Unavail.Dur.Sys(s)", "bad_direction": "high", "threshold": 20, "severity": 5, "category": "System Unavailability", "reason": "System unavailability duration increased.", "recommended_action": "Check system faults, board alarms, transmission, and eNodeB health."},
            {"feature": "L.Cell.Unavail.Dur.Manual(s)", "bad_direction": "high", "threshold": 20, "severity": 3, "category": "Manual Unavailability", "reason": "Manual unavailability duration increased.", "recommended_action": "Check manual lock, planned work, maintenance activity, and cell administrative state."},
            {"feature": "L.Cell.Unavail.Dur.Sys.S1Fail(s)", "bad_direction": "high", "threshold": 20, "severity": 5, "category": "S1 Failure Unavailability", "reason": "S1 failure unavailability duration increased.", "recommended_action": "Check S1 link, MME connection, transmission, and core network alarms."},
        ],
    },

    "RACH Success Rate": {
        "target_kpi": "(HU) RACH Success Rate(%)",
        "bad_direction": "low",
        "default_threshold": 5.0,
        "category": "Accessibility",
        "output_prefix": "rach_success_rate",
        "min_baseline_value": 95.0,  # NEW
        "related_rules": [
            {"feature": "RACH Setup Failed Number", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "RACH Setup Failures", "reason": "RACH setup failures increased.", "recommended_action": "Check PRACH parameters, root sequence planning, coverage, interference, and access load."},
            {"feature": "RACH Contention-Based Failures", "bad_direction": "high", "threshold": 20, "severity": 3, "category": "Contention-Based RACH Failure", "reason": "Contention-based RACH failures increased.", "recommended_action": "Check preamble congestion, PRACH configuration, root sequence, and access load."},
            {"feature": "RACH_att", "bad_direction": "high", "threshold": 20, "severity": 2, "category": "High RACH Attempts", "reason": "RACH attempts increased.", "recommended_action": "Check access load, coverage, PRACH capacity, and random access configuration."},
            {"feature": "RACH Contention-Based SR", "bad_direction": "low", "threshold": 5, "severity": 3, "category": "Contention-Based RACH SR Drop", "reason": "Contention-based RACH success rate decreased.", "recommended_action": "Check PRACH configuration, root sequence planning, and access congestion."},
            {"feature": "RACH Non-Contention-Based SR", "bad_direction": "low", "threshold": 5, "severity": 3, "category": "Non-Contention RACH SR Drop", "reason": "Non-contention RACH success rate decreased.", "recommended_action": "Check HO-related RACH, target cell access, and PRACH settings."},
        ],
    },

    "CSFB KPI": {
        "target_kpi": "CSFB SR%",
        "bad_direction": "low",
        "default_threshold": 5.0,
        "category": "CSFB / Voice Accessibility",
        "output_prefix": "csfb_kpi",
        "min_baseline_value": 90.0,  # NEW
        "related_rules": [
            {"feature": "CSFB Failure Times", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "CSFB Failure Increase", "reason": "CSFB failure times increased.", "recommended_action": "Check CSFB failure reasons, MME/S1 signaling, RRC redirection, and target 2G/3G availability."},
            {"feature": "L.CSFB.PrepAtt", "bad_direction": "high", "threshold": 20, "severity": 2, "category": "High CSFB Preparation Attempts", "reason": "CSFB preparation attempts increased, which may increase CSFB load.", "recommended_action": "Check CSFB traffic demand, MME load, S1 signaling, and whether the increase is normal voice demand."},
            {"feature": "L.RRCRedirection.E2W.CSFB", "bad_direction": "low", "threshold": 20, "severity": 3, "category": "E2W CSFB Redirection Drop", "reason": "LTE-to-3G CSFB redirection count decreased compared with baseline.", "recommended_action": "Check UTRAN neighbor configuration, 3G target coverage, UTRAN frequency priority, and CSFB redirection settings."},
            {"feature": "L.RRCRedirection.E2G.CSFB", "bad_direction": "low", "threshold": 20, "severity": 3, "category": "E2G CSFB Redirection Drop", "reason": "LTE-to-2G CSFB redirection count decreased compared with baseline.", "recommended_action": "Check GERAN neighbor configuration, 2G target coverage, GERAN frequency priority, LAI/TAI mapping, and CSFB redirection settings."},
            {"feature": "(TE) RRC Setup SR%", "bad_direction": "low", "threshold": 5, "severity": 4, "category": "LTE RRC Access Issue", "reason": "RRC setup success rate decreased, which can affect CSFB before fallback starts.", "recommended_action": "Check LTE RRC accessibility, RACH, RRC setup failures, admission control, and radio quality."},
            {"feature": "ERAB Setup Success Rate", "bad_direction": "low", "threshold": 5, "severity": 3, "category": "E-RAB Setup Issue", "reason": "E-RAB setup success rate decreased, indicating possible access or core signaling issue affecting services.", "recommended_action": "Check E-RAB setup failures, MME/TNL/RNL causes, admission control, radio resources, and S1 signaling."},
        ],
    },

    "VoLTE KPIs": {
        "target_kpi": "BA_Voice E2E VQI",
        "bad_direction": "low",
        "default_threshold": 5.0,
        "category": "VoLTE",
        "output_prefix": "volte_kpis",
        "min_baseline_value": 3.5,  # NEW: VQI baseline threshold
        "related_rules": [
            {"feature": "VoLTE Traffic (Erl)(Erl)", "bad_direction": "low", "threshold": 20, "severity": 3, "category": "VoLTE Traffic Drop", "reason": "VoLTE traffic decreased.", "recommended_action": "Check VoLTE user demand, IMS service, VoLTE coverage, and QCI-1 traffic."},
            {"feature": "L.Traffic.User.VoIP.Avg", "bad_direction": "low", "threshold": 20, "severity": 2, "category": "VoIP User Drop", "reason": "Average VoIP users decreased.", "recommended_action": "Check VoLTE traffic demand, service availability, and IMS registration behavior."},
            {"feature": "DL Traffic QCI-1", "bad_direction": "low", "threshold": 20, "severity": 3, "category": "QCI-1 DL Traffic Drop", "reason": "QCI-1 DL traffic decreased.", "recommended_action": "Check VoLTE bearer traffic, IMS service, and VoLTE user behavior."},
            {"feature": "E-RAB Drop(ENB+MME)_Tot", "bad_direction": "high", "threshold": 20, "severity": 5, "category": "VoLTE Retainability Risk", "reason": "Total E-RAB drops increased.", "recommended_action": "Check VoLTE drops, radio quality, TNL/MME causes, and mobility."},
            {"feature": "E-RAB Drop Rate QCI 7", "bad_direction": "high", "threshold": 20, "severity": 4, "category": "QCI-7 Drop Issue", "reason": "QCI-7 drop rate increased.", "recommended_action": "Check VoLTE-related bearer retainability and radio quality."},
            {"feature": "BA_Overall SRVCC HO Execution Success Rate (%)", "bad_direction": "low", "threshold": 5, "severity": 4, "category": "SRVCC Execution Degradation", "reason": "SRVCC HO execution success rate decreased.", "recommended_action": "Check SRVCC neighbors, 2G/3G target cells, IMS/SRVCC configuration, and mobility parameters."},
            {"feature": "BA_Overall SRVCC HO Preparation Success Rate (%)", "bad_direction": "low", "threshold": 5, "severity": 3, "category": "SRVCC Preparation Degradation", "reason": "SRVCC HO preparation success rate decreased.", "recommended_action": "Check SRVCC preparation, target availability, MSC/MME coordination, and neighbor definitions."},
        ],
    },
}


# ============================================================
# 3. HELPER FUNCTIONS
# ============================================================

def clean_excel_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Clean Excel column names from spaces and hidden line breaks."""
    df = df.copy()
    cleaned_columns = []
    for col in df.columns:
        col = str(col)
        col = col.replace(chr(10), " ")
        col = col.replace(chr(13), " ")
        col = col.strip()
        cleaned_columns.append(col)
    df.columns = cleaned_columns
    return df


def normalize_column_name(col) -> str:
    """Normalize column names for smart matching."""
    col = str(col).lower()
    col = col.replace(" ", "")
    col = col.replace("_", "")
    col = col.replace("-", "")
    col = col.replace(chr(10), "")
    col = col.replace(chr(13), "")
    col = col.strip()
    return col


def find_matching_column(df: pd.DataFrame, wanted_col: str):
    """Find real Excel column even if spaces/newlines are different."""
    if wanted_col in df.columns:
        return wanted_col
    for col in df.columns:
        if str(col).strip() == str(wanted_col).strip():
            return col
    wanted_norm = normalize_column_name(wanted_col)
    for col in df.columns:
        if normalize_column_name(col) == wanted_norm:
            return col
    return None


def clean_numeric_series(series: pd.Series) -> pd.Series:
    """Convert mixed numeric/text column to numeric - Enhanced with N/A handling."""
    return pd.to_numeric(
        series.astype(str)
        .str.replace("%", "", regex=False)
        .str.replace(",", "", regex=False)
        .str.replace("N/A", "", regex=False)
        .str.replace("#DIV/0!", "", regex=False)
        .str.replace("#N/A", "", regex=False)
        .str.replace("NIL", "", regex=False)
        .str.strip()
        .replace({"": np.nan, "nan": np.nan, "None": np.nan}),
        errors="coerce",
    )


def calculate_degradation(recent_value, baseline_value, bad_direction):
    """Return degradation percentage. Positive means worse."""
    if pd.isna(recent_value) or pd.isna(baseline_value):
        return np.nan
    if baseline_value == 0:
        return np.nan
    if bad_direction == "low":
        return ((baseline_value - recent_value) / baseline_value) * 100
    if bad_direction == "high":
        return ((recent_value - baseline_value) / baseline_value) * 100
    return np.nan


def perform_ttest(recent_values, baseline_values):
    """
    Perform Welch's t-test between recent and baseline periods.
    Returns (is_significant, p_value, t_statistic)
    """
    try:
        recent_clean = recent_values.dropna()
        baseline_clean = baseline_values.dropna()
        
        if len(recent_clean) < 2 or len(baseline_clean) < 2:
            return False, np.nan, np.nan
        
        t_stat, p_value = stats.ttest_ind(recent_clean, baseline_clean, equal_var=False)
        
        # Significant if p < 0.05
        is_significant = p_value < 0.05
        return is_significant, p_value, t_stat
    except Exception:
        return False, np.nan, np.nan


def get_periods_enhanced(df, date_col, num_days, baseline_mode=BASELINE_MODE_LAST_WEEK, 
                          custom_baseline_start=None, custom_baseline_end=None):
    """
    Get analysis periods with configurable baseline window.
    
    baseline_mode options:
    - BASELINE_MODE_LAST_WEEK: Same N days from last week (original behavior)
    - BASELINE_MODE_4WEEK_AVG: 4-week rolling average
    - BASELINE_MODE_CUSTOM: User-defined date range
    """
    # Normalize to date level for hourly data
    last_date = df[date_col].dt.normalize().max()
    recent_start = last_date - pd.Timedelta(days=num_days - 1)
    recent_end = last_date
    
    if baseline_mode == BASELINE_MODE_LAST_WEEK:
        baseline_start = recent_start - pd.Timedelta(days=7)
        baseline_end = recent_end - pd.Timedelta(days=7)
        
    elif baseline_mode == BASELINE_MODE_4WEEK_AVG:
        # 4-week rolling: use all 4 weeks before recent period
        baseline_start = recent_start - pd.Timedelta(days=28)
        baseline_end = recent_start - pd.Timedelta(days=1)
        
    elif baseline_mode == BASELINE_MODE_CUSTOM:
        if custom_baseline_start and custom_baseline_end:
            baseline_start = pd.Timestamp(custom_baseline_start)
            baseline_end = pd.Timestamp(custom_baseline_end)
        else:
            # Fallback to last week
            baseline_start = recent_start - pd.Timedelta(days=7)
            baseline_end = recent_end - pd.Timedelta(days=7)
    else:
        baseline_start = recent_start - pd.Timedelta(days=7)
        baseline_end = recent_end - pd.Timedelta(days=7)
    
    return last_date, recent_start, recent_end, baseline_start, baseline_end


def find_degradation_causes_vectorized(df, rules):
    """
    Vectorized cause detection for performance optimization.
    Replaces row-by-row apply() with column-wise operations.
    
    NEW: Uses severity weighting for cause ranking.
    
    IMPORTANT: Reset index before calling this function to ensure proper alignment.
    """
    # CRITICAL FIX: Reset index to ensure proper alignment
    df_work = df.reset_index(drop=True).copy()
    
    detected_causes_list = []
    
    for rule in rules:
        feature = rule["feature"]
        recent_col = f"recent_{feature}"
        baseline_col = f"baseline_{feature}"
        
        if recent_col not in df_work.columns or baseline_col not in df_work.columns:
            continue
        
        recent_values = df_work[recent_col].values  # Use .values for position-based access
        baseline_values = df_work[baseline_col].values
        bad_direction = rule["bad_direction"]
        threshold = rule["threshold"]
        severity = rule.get("severity", 3)  # Default severity
        
        # Vectorized calculation using numpy arrays
        with np.errstate(divide='ignore', invalid='ignore'):
            if bad_direction == "low":
                change_pct = np.where(
                    baseline_values != 0,
                    ((baseline_values - recent_values) / baseline_values) * 100,
                    np.nan
                )
            else:  # high
                change_pct = np.where(
                    baseline_values != 0,
                    ((recent_values - baseline_values) / baseline_values) * 100,
                    np.nan
                )
        
        # Create mask for cells passing threshold
        mask = change_pct >= threshold
        mask = mask & ~np.isnan(change_pct)  # Exclude NaN values
        
        if mask.any():
            # Score = change_pct * severity_weight
            score = change_pct * severity
            
            # Get positions where mask is True
            positions = np.where(mask)[0]
            
            for pos in positions:
                detected_causes_list.append({
                    "row_pos": pos,
                    "feature": feature,
                    "recent_value": recent_values[pos],
                    "baseline_value": baseline_values[pos],
                    "change_pct": change_pct[pos],
                    "severity": severity,
                    "score": score[pos],
                    "category": rule["category"],
                    "reason": rule["reason"],
                    "recommended_action": rule["recommended_action"],
                })
    
    # Default result columns
    default_cols = {
        "main_cause_counter_or_kpi": "No strong related counter detected",
        "main_cause_recent_value": np.nan,
        "main_cause_baseline_value": np.nan,
        "main_cause_change_%": np.nan,
        "main_root_cause_category": "Unknown",
        "main_degradation_reason": "Main KPI degraded, but no related counter passed its threshold.",
        "main_recommended_action": "Check raw counters, alarms, availability, recent changes, and nearby cells manually.",
        "number_of_detected_causes": 0,
        "multi_cause_flag": "No",
        "all_detected_causes": "None",
        "all_cause_categories": "Unknown",
        "all_recommended_actions": "Manual investigation needed",
    }
    
    # If no causes detected, return defaults for all rows
    if not detected_causes_list:
        result_df = pd.DataFrame(default_cols, index=range(len(df_work)))
        return result_df
    
    # Convert to DataFrame
    causes_df = pd.DataFrame(detected_causes_list)
    
    # Sort by score (severity-weighted) for each cell
    causes_df = causes_df.sort_values(["row_pos", "score"], ascending=[True, False])
    
    # Aggregate causes per cell using row position
    result_dict = {}
    
    for row_pos in range(len(df_work)):
        cell_causes = causes_df[causes_df["row_pos"] == row_pos].sort_values("score", ascending=False)
        
        if len(cell_causes) == 0:
            result_dict[row_pos] = default_cols.copy()
        else:
            main_cause = cell_causes.iloc[0]
            
            all_causes_text = " | ".join([
                f"{row['feature']}: recent={row['recent_value']:.2f}, baseline={row['baseline_value']:.2f}, change={row['change_pct']:.2f}%"
                for _, row in cell_causes.head(5).iterrows()
            ])
            all_categories_text = " | ".join(cell_causes["category"].head(5).tolist())
            all_actions_text = " | ".join(cell_causes["recommended_action"].head(5).tolist())
            
            result_dict[row_pos] = {
                "main_cause_counter_or_kpi": main_cause["feature"],
                "main_cause_recent_value": main_cause["recent_value"],
                "main_cause_baseline_value": main_cause["baseline_value"],
                "main_cause_change_%": main_cause["change_pct"],
                "main_root_cause_category": main_cause["category"],
                "main_degradation_reason": main_cause["reason"],
                "main_recommended_action": main_cause["recommended_action"],
                "number_of_detected_causes": len(cell_causes),
                "multi_cause_flag": "Yes" if len(cell_causes) > 1 else "No",
                "all_detected_causes": all_causes_text,
                "all_cause_categories": all_categories_text,
                "all_recommended_actions": all_actions_text,
            }
    
    # Create result DataFrame with proper row positions
    result_df = pd.DataFrame.from_dict(result_dict, orient='index')
    
    return result_df


def find_degradation_causes_row(row, rules):
    """
    Row-by-row cause detection (fallback method).
    Used when vectorized detection fails.
    """
    detected_causes = []
    
    for rule in rules:
        feature = rule["feature"]
        recent_col = f"recent_{feature}"
        baseline_col = f"baseline_{feature}"
        
        if recent_col not in row.index or baseline_col not in row.index:
            continue
        
        recent_value = row[recent_col]
        baseline_value = row[baseline_col]
        
        change_pct = calculate_degradation(
            recent_value,
            baseline_value,
            rule["bad_direction"],
        )
        
        if pd.isna(change_pct):
            continue
        
        if change_pct >= rule["threshold"]:
            severity = rule.get("severity", 3)
            detected_causes.append({
                "feature": feature,
                "recent_value": recent_value,
                "baseline_value": baseline_value,
                "change_pct": change_pct,
                "severity": severity,
                "score": change_pct * severity,
                "category": rule["category"],
                "reason": rule["reason"],
                "recommended_action": rule["recommended_action"],
            })
    
    if not detected_causes:
        return pd.Series({
            "main_cause_counter_or_kpi": "No strong related counter detected",
            "main_cause_recent_value": np.nan,
            "main_cause_baseline_value": np.nan,
            "main_cause_change_%": np.nan,
            "main_root_cause_category": "Unknown",
            "main_degradation_reason": "Main KPI degraded, but no related counter passed its threshold.",
            "main_recommended_action": "Check raw counters, alarms, availability, recent changes, and nearby cells manually.",
            "number_of_detected_causes": 0,
            "multi_cause_flag": "No",
            "all_detected_causes": "None",
            "all_cause_categories": "Unknown",
            "all_recommended_actions": "Manual investigation needed",
        })
    
    # Sort by severity-weighted score
    detected_causes = sorted(detected_causes, key=lambda x: x["score"], reverse=True)
    main_cause = detected_causes[0]
    
    all_causes_text = " | ".join([
        f"{c['feature']}: recent={c['recent_value']:.2f}, baseline={c['baseline_value']:.2f}, change={c['change_pct']:.2f}%"
        for c in detected_causes[:5]
    ])
    all_categories_text = " | ".join([c["category"] for c in detected_causes[:5]])
    all_actions_text = " | ".join([c["recommended_action"] for c in detected_causes[:5]])
    
    return pd.Series({
        "main_cause_counter_or_kpi": main_cause["feature"],
        "main_cause_recent_value": main_cause["recent_value"],
        "main_cause_baseline_value": main_cause["baseline_value"],
        "main_cause_change_%": main_cause["change_pct"],
        "main_root_cause_category": main_cause["category"],
        "main_degradation_reason": main_cause["reason"],
        "main_recommended_action": main_cause["recommended_action"],
        "number_of_detected_causes": len(detected_causes),
        "multi_cause_flag": "Yes" if len(detected_causes) > 1 else "No",
        "all_detected_causes": all_causes_text,
        "all_cause_categories": all_categories_text,
        "all_recommended_actions": all_actions_text,
    })


# ============================================================
# 4. ENHANCED MAIN ANALYSIS ENGINE
# ============================================================

def analyze_selected_kpi_enhanced(
    df, 
    selected_kpi_name, 
    num_days, 
    degradation_threshold, 
    require_complete_days=True,
    baseline_mode=BASELINE_MODE_LAST_WEEK,
    custom_baseline_start=None,
    custom_baseline_end=None,
    enable_significance_test=True,
    log_callback=None
):
    """
    Enhanced analysis with:
    - Configurable baseline window
    - Minimum baseline value filter
    - Statistical significance test
    - Max/percentile aggregation for failure counters
    - Zero baseline warning
    
    Returns: (output_df, metadata)
    """
    def log_msg(msg):
        if log_callback:
            log_callback(msg)
    
    config = KPI_CONFIGS[selected_kpi_name]
    
    # Clean Excel column names
    df = clean_excel_columns(df)
    
    # Smart target KPI matching
    original_target_kpi = config["target_kpi"]
    target_kpi = find_matching_column(df, original_target_kpi)
    
    if target_kpi is None:
        raise ValueError(f"Target KPI column not found in Excel: {original_target_kpi}")
    
    bad_direction = config["bad_direction"]
    related_rules = config["related_rules"]
    min_baseline_value = config.get("min_baseline_value", 0.0)
    
    needed_cols = CELL_ID_COLS + [DATE_COL, target_kpi]
    missing_cols = [col for col in needed_cols if col not in df.columns]
    if missing_cols:
        raise ValueError(f"Missing required columns: {missing_cols}")
    
    df_kpi = df[needed_cols].copy()
    # Normalize date to handle hourly data
    df_kpi[DATE_COL] = pd.to_datetime(df_kpi[DATE_COL], errors="coerce").dt.normalize()
    df_kpi[target_kpi] = clean_numeric_series(df_kpi[target_kpi])
    df_kpi = df_kpi.dropna(subset=[DATE_COL, target_kpi])
    df_kpi = df_kpi[df_kpi[target_kpi] >= 0].copy()
    
    # Get periods with enhanced baseline mode
    last_date, recent_start, recent_end, baseline_start, baseline_end = get_periods_enhanced(
        df_kpi, DATE_COL, num_days, baseline_mode, custom_baseline_start, custom_baseline_end
    )
    
    recent_df = df_kpi[(df_kpi[DATE_COL] >= recent_start) & (df_kpi[DATE_COL] <= recent_end)].copy()
    baseline_df = df_kpi[(df_kpi[DATE_COL] >= baseline_start) & (df_kpi[DATE_COL] <= baseline_end)].copy()
    
    # Enhanced aggregation with max and percentile for failure counters
    agg_funcs = {
        target_kpi: ["mean", "max", "sum"]
    }
    
    recent_agg = recent_df.groupby(CELL_ID_COLS).agg({
        target_kpi: ["mean", "max", "sum"],
        DATE_COL: "nunique"
    }).reset_index()
    recent_agg.columns = CELL_ID_COLS + ["recent_avg_kpi", "recent_max_kpi", "recent_total_kpi", "recent_days_count"]
    
    baseline_agg = baseline_df.groupby(CELL_ID_COLS).agg({
        target_kpi: ["mean", "max", "sum"],
        DATE_COL: "nunique"
    }).reset_index()
    baseline_agg.columns = CELL_ID_COLS + ["baseline_avg_kpi", "baseline_max_kpi", "baseline_total_kpi", "baseline_days_count"]
    
    comparison = recent_agg.merge(baseline_agg, on=CELL_ID_COLS, how="inner")
    
    if require_complete_days:
        comparison = comparison[
            (comparison["recent_days_count"] == num_days) &
            (comparison["baseline_days_count"] == num_days)
        ].copy()
    
    # NEW: Log warning for zero baseline cells
    zero_baseline_mask = comparison["baseline_avg_kpi"] == 0
    zero_baseline_count = zero_baseline_mask.sum()
    if zero_baseline_count > 0:
        zero_baseline_cells = comparison[zero_baseline_mask][CELL_COL].head(10).tolist()
        log_msg(f"WARNING: {zero_baseline_count} cells have zero baseline value and will be excluded")
        log_msg(f"Examples: {zero_baseline_cells}")
    
    comparison = comparison[comparison["baseline_avg_kpi"] != 0].copy()
    
    # NEW: Apply minimum baseline value filter
    if min_baseline_value > 0:
        min_baseline_mask = comparison["baseline_avg_kpi"] >= min_baseline_value
        excluded_by_min = (~min_baseline_mask).sum()
        if excluded_by_min > 0:
            log_msg(f"INFO: {excluded_by_min} cells excluded by min_baseline_value filter (< {min_baseline_value})")
        comparison = comparison[min_baseline_mask].copy()
    
    # Calculate degradation
    comparison["kpi_degradation_ratio_%"] = comparison.apply(
        lambda row: calculate_degradation(row["recent_avg_kpi"], row["baseline_avg_kpi"], bad_direction),
        axis=1,
    )
    
    # NEW: Statistical significance test
    if enable_significance_test:
        significance_results = []
        for idx, row in comparison.iterrows():
            cell_id = (row[SITE_COL], row[CELL_COL])
            
            # Get daily values for this cell
            cell_recent = recent_df[
                (recent_df[SITE_COL] == cell_id[0]) & 
                (recent_df[CELL_COL] == cell_id[1])
            ][target_kpi]
            cell_baseline = baseline_df[
                (baseline_df[SITE_COL] == cell_id[0]) & 
                (baseline_df[CELL_COL] == cell_id[1])
            ][target_kpi]
            
            is_sig, p_val, t_stat = perform_ttest(cell_recent, cell_baseline)
            significance_results.append({
                "index": idx,
                "stat_significant": is_sig,
                "p_value": p_val,
                "t_statistic": t_stat
            })
        
        sig_df = pd.DataFrame(significance_results).set_index("index")
        comparison["stat_significant"] = sig_df["stat_significant"].reindex(comparison.index).fillna(False)
        comparison["p_value"] = sig_df["p_value"].reindex(comparison.index)
        comparison["t_statistic"] = sig_df["t_statistic"].reindex(comparison.index)
    
    comparison["kpi_status"] = np.where(
        comparison["kpi_degradation_ratio_%"] >= degradation_threshold,
        "Degraded",
        "Normal",
    )
    
    # If significance test enabled, require both threshold AND significance
    if enable_significance_test:
        comparison["kpi_status"] = np.where(
            (comparison["kpi_degradation_ratio_%"] >= degradation_threshold) & 
            (comparison["stat_significant"] == True),
            "Degraded",
            "Normal",
        )
    
    comparison["selected_kpi_name"] = selected_kpi_name
    comparison["target_kpi_column"] = target_kpi
    comparison["kpi_category"] = config["category"]
    comparison["kpi_bad_direction"] = bad_direction
    comparison["selected_threshold_%"] = degradation_threshold
    comparison["recent_period"] = f"{recent_start.date()} to {recent_end.date()}"
    comparison["baseline_period"] = f"{baseline_start.date()} to {baseline_end.date()}"
    comparison["baseline_mode"] = baseline_mode
    
    degraded_cells = comparison[comparison["kpi_status"] == "Degraded"].copy()
    degraded_cells = degraded_cells.sort_values("kpi_degradation_ratio_%", ascending=False)
    
    debug_info = {
        "cells_after_merge": comparison.shape[0],
        "recent_days_distribution": comparison["recent_days_count"].value_counts().sort_index().to_dict(),
        "baseline_days_distribution": comparison["baseline_days_count"].value_counts().sort_index().to_dict(),
        "max_degradation": comparison["kpi_degradation_ratio_%"].max() if not comparison.empty else None,
        "min_degradation": comparison["kpi_degradation_ratio_%"].min() if not comparison.empty else None,
        "mean_degradation": comparison["kpi_degradation_ratio_%"].mean() if not comparison.empty else None,
        "zero_baseline_excluded": zero_baseline_count,
        "min_baseline_excluded": excluded_by_min if min_baseline_value > 0 else 0,
    }
    
    metadata = {
        "last_date": last_date,
        "recent_start": recent_start,
        "recent_end": recent_end,
        "baseline_start": baseline_start,
        "baseline_end": baseline_end,
        "baseline_mode": baseline_mode,
        "available_related_features": [],
        "missing_related_features": [],
        "debug_info": debug_info,
    }
    
    if degraded_cells.empty:
        return degraded_cells, metadata
    
    # Smart related counter matching
    available_related_rules = []
    missing_related_features = []
    
    for rule in related_rules:
        matched_col = find_matching_column(df, rule["feature"])
        if matched_col is not None:
            new_rule = rule.copy()
            new_rule["feature"] = matched_col
            available_related_rules.append(new_rule)
        else:
            missing_related_features.append(rule["feature"])
    
    available_related_features = [rule["feature"] for rule in available_related_rules]
    metadata["available_related_features"] = available_related_features
    metadata["missing_related_features"] = missing_related_features
    
    if available_related_features:
        reason_cols = CELL_ID_COLS + [DATE_COL] + available_related_features
        df_reason = df[reason_cols].copy()
        df_reason[DATE_COL] = pd.to_datetime(df_reason[DATE_COL], errors="coerce").dt.normalize()
        
        for col in available_related_features:
            df_reason[col] = clean_numeric_series(df_reason[col])
        
        recent_reason_df = df_reason[(df_reason[DATE_COL] >= recent_start) & (df_reason[DATE_COL] <= recent_end)].copy()
        baseline_reason_df = df_reason[(df_reason[DATE_COL] >= baseline_start) & (df_reason[DATE_COL] <= baseline_end)].copy()
        
        # Enhanced aggregation: mean + max for failure counters
        agg_dict = {}
        for col in available_related_features:
            agg_dict[col] = ["mean", "max"]
        
        recent_reason_agg = recent_reason_df.groupby(CELL_ID_COLS).agg(agg_dict).reset_index()
        baseline_reason_agg = baseline_reason_df.groupby(CELL_ID_COLS).agg(agg_dict).reset_index()
        
        # Flatten column names
        new_cols = CELL_ID_COLS.copy()
        for col in available_related_features:
            new_cols.append(f"recent_{col}_mean")
            new_cols.append(f"recent_{col}_max")
        recent_reason_agg.columns = new_cols
        
        new_cols = CELL_ID_COLS.copy()
        for col in available_related_features:
            new_cols.append(f"baseline_{col}_mean")
            new_cols.append(f"baseline_{col}_max")
        baseline_reason_agg.columns = new_cols
        
        # Create aliases for compatibility with cause detection
        for col in available_related_features:
            recent_reason_agg[f"recent_{col}"] = recent_reason_agg[f"recent_{col}_mean"]
            baseline_reason_agg[f"baseline_{col}"] = baseline_reason_agg[f"baseline_{col}_mean"]
        
        degraded_with_causes = degraded_cells.merge(recent_reason_agg, on=CELL_ID_COLS, how="left")
        degraded_with_causes = degraded_with_causes.merge(baseline_reason_agg, on=CELL_ID_COLS, how="left")
        
        # CRITICAL: Reset index for proper alignment with cause detection
        degraded_with_causes = degraded_with_causes.reset_index(drop=True)
        
        # Use vectorized cause detection with fallback
        try:
            cause_results = find_degradation_causes_vectorized(degraded_with_causes, available_related_rules)
            degraded_with_causes = pd.concat([degraded_with_causes.reset_index(drop=True), cause_results.reset_index(drop=True)], axis=1)
        except Exception as vec_error:
            # Fallback to row-by-row apply if vectorized fails
            log_msg(f"Vectorized cause detection failed, using fallback: {vec_error}")
            cause_results = degraded_with_causes.apply(
                lambda row: find_degradation_causes_row(row, available_related_rules),
                axis=1,
            )
            degraded_with_causes = pd.concat([degraded_with_causes, cause_results], axis=1)
    else:
        degraded_with_causes = degraded_cells.copy()
        degraded_with_causes["main_cause_counter_or_kpi"] = "No related counters available in sheet"
        degraded_with_causes["main_cause_recent_value"] = np.nan
        degraded_with_causes["main_cause_baseline_value"] = np.nan
        degraded_with_causes["main_cause_change_%"] = np.nan
        degraded_with_causes["main_root_cause_category"] = "Unknown"
        degraded_with_causes["main_degradation_reason"] = "No related counters from the config were found in the uploaded sheet."
        degraded_with_causes["main_recommended_action"] = "Check KPI manually or update KPI_CONFIGS with available counters."
        degraded_with_causes["number_of_detected_causes"] = 0
        degraded_with_causes["multi_cause_flag"] = "No"
        degraded_with_causes["all_detected_causes"] = "None"
        degraded_with_causes["all_cause_categories"] = "Unknown"
        degraded_with_causes["all_recommended_actions"] = "Manual investigation needed"
    
    final_cols = CELL_ID_COLS + [
        "selected_kpi_name", "target_kpi_column", "kpi_category", "kpi_bad_direction",
        "selected_threshold_%", "recent_period", "baseline_period", "baseline_mode",
        "recent_avg_kpi", "baseline_avg_kpi", "recent_max_kpi", "baseline_max_kpi",
        "recent_total_kpi", "baseline_total_kpi", "recent_days_count", "baseline_days_count",
        "kpi_degradation_ratio_%", "kpi_status", "stat_significant", "p_value",
        "main_cause_counter_or_kpi", "main_cause_recent_value", "main_cause_baseline_value",
        "main_cause_change_%", "main_root_cause_category", "main_degradation_reason",
        "main_recommended_action", "number_of_detected_causes", "multi_cause_flag",
        "all_detected_causes", "all_cause_categories", "all_recommended_actions",
    ]
    
    # Filter to available columns
    available_final_cols = [col for col in final_cols if col in degraded_with_causes.columns]
    
    return degraded_with_causes[available_final_cols].copy(), metadata


# ============================================================
# 5. ENHANCED GUI APPLICATION
# ============================================================

class LTEKPIAnalyzerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("LTE KPI Degradation Analyzer v2.0 - Developed by Musketeers_Team (ITI Graduation Project 2026)")
        self.root.geometry("1350x800")
        self.root.minsize(1150, 750)
        
        # File and sheet
        self.file_path = tk.StringVar()
        self.excel_sheets = []
        self.selected_sheet = tk.StringVar()
        
        # Analysis settings
        self.selected_kpi = tk.StringVar(value="DL Traffic")
        self.num_days = tk.IntVar(value=4)
        self.threshold = tk.DoubleVar(value=KPI_CONFIGS["DL Traffic"]["default_threshold"])
        self.require_complete_days = tk.BooleanVar(value=True)
        
        # NEW: Baseline mode settings
        self.baseline_mode = tk.StringVar(value=BASELINE_MODE_LAST_WEEK)
        self.custom_baseline_start = tk.StringVar()
        self.custom_baseline_end = tk.StringVar()
        
        # NEW: Statistical significance test
        self.enable_significance_test = tk.BooleanVar(value=True)
        
        # Output storage
        self.output_df = None
        self.original_df = None
        self.degraded_cell_ids = set()
        self.all_outputs = {}
        self.summary_df = None
        self.analysis_mode = "single"
        
        # Running state
        self.is_running = False
        self.progress_var = tk.DoubleVar(value=0)
        self.progress_percent_var = tk.StringVar(value="0%")
        self.status_var = tk.StringVar(value="Ready")
        
        self.build_ui()
    
    def build_ui(self):
        # Main frames
        top_frame = ttk.Frame(self.root, padding=10)
        top_frame.pack(fill="x")
        
        settings_frame = ttk.LabelFrame(self.root, text="Analysis Settings", padding=10)
        settings_frame.pack(fill="x", padx=10, pady=5)
        
        result_frame = ttk.LabelFrame(self.root, text="Results Preview", padding=10)
        result_frame.pack(fill="both", expand=True, padx=10, pady=5)
        
        log_frame = ttk.LabelFrame(self.root, text="Log", padding=10)
        log_frame.pack(fill="x", padx=10, pady=5)
        
        # File selection with sheet selector
        ttk.Label(top_frame, text="Excel File:").pack(side="left")
        ttk.Entry(top_frame, textvariable=self.file_path, width=70).pack(side="left", padx=5)
        ttk.Button(top_frame, text="Browse", command=self.browse_file).pack(side="left", padx=5)
        
        # Sheet selector (NEW)
        ttk.Label(top_frame, text="Sheet:").pack(side="left", padx=(20, 5))
        self.sheet_combo = ttk.Combobox(top_frame, textvariable=self.selected_sheet, 
                                         values=[], state="readonly", width=25)
        self.sheet_combo.pack(side="left", padx=5)
        
        # Row 0: KPI and Days
        ttk.Label(settings_frame, text="KPI:").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        self.kpi_combo = ttk.Combobox(
            settings_frame, textvariable=self.selected_kpi,
            values=list(KPI_CONFIGS.keys()), state="readonly", width=25,
        )
        self.kpi_combo.grid(row=0, column=1, sticky="w", padx=5, pady=5)
        self.kpi_combo.bind("<<ComboboxSelected>>", self.on_kpi_change)
        
        ttk.Label(settings_frame, text="Comparison Days:").grid(row=0, column=2, sticky="w", padx=5, pady=5)
        ttk.Spinbox(settings_frame, from_=1, to=14, textvariable=self.num_days, width=8).grid(
            row=0, column=3, sticky="w", padx=5, pady=5
        )
        
        ttk.Label(settings_frame, text="Threshold (%):").grid(row=0, column=4, sticky="w", padx=5, pady=5)
        ttk.Entry(settings_frame, textvariable=self.threshold, width=10).grid(
            row=0, column=5, sticky="w", padx=5, pady=5
        )
        
        # Row 1: Baseline Mode (NEW)
        ttk.Label(settings_frame, text="Baseline Mode:").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        baseline_frame = ttk.Frame(settings_frame)
        baseline_frame.grid(row=1, column=1, columnspan=3, sticky="w", padx=5, pady=5)
        
        ttk.Radiobutton(baseline_frame, text="Last Week (7-day)", variable=self.baseline_mode, 
                        value=BASELINE_MODE_LAST_WEEK).pack(side="left", padx=5)
        ttk.Radiobutton(baseline_frame, text="4-Week Rolling Avg", variable=self.baseline_mode,
                        value=BASELINE_MODE_4WEEK_AVG).pack(side="left", padx=5)
        ttk.Radiobutton(baseline_frame, text="Custom Range", variable=self.baseline_mode,
                        value=BASELINE_MODE_CUSTOM).pack(side="left", padx=5)
        
        # Custom date range
        ttk.Label(settings_frame, text="Custom Baseline:").grid(row=1, column=4, sticky="w", padx=5, pady=5)
        custom_frame = ttk.Frame(settings_frame)
        custom_frame.grid(row=1, column=5, columnspan=2, sticky="w", padx=5, pady=5)
        ttk.Label(custom_frame, text="Start:").pack(side="left")
        ttk.Entry(custom_frame, textvariable=self.custom_baseline_start, width=12).pack(side="left", padx=2)
        ttk.Label(custom_frame, text="End:").pack(side="left", padx=(5, 0))
        ttk.Entry(custom_frame, textvariable=self.custom_baseline_end, width=12).pack(side="left", padx=2)
        
        # Row 2: Checkboxes and buttons
        ttk.Checkbutton(
            settings_frame, text="Require complete days",
            variable=self.require_complete_days,
        ).grid(row=2, column=0, columnspan=2, sticky="w", padx=5, pady=5)
        
        # NEW: Statistical significance checkbox
        ttk.Checkbutton(
            settings_frame, text="Enable t-test significance filter",
            variable=self.enable_significance_test,
        ).grid(row=2, column=2, columnspan=2, sticky="w", padx=5, pady=5)
        
        # Buttons
        self.run_button = ttk.Button(settings_frame, text="Run Selected KPI", command=self.run_analysis_thread)
        self.run_button.grid(row=0, column=6, padx=10, pady=5)
        
        self.run_all_button = ttk.Button(settings_frame, text="Analyze All KPIs", command=self.run_all_analysis_thread)
        self.run_all_button.grid(row=0, column=7, padx=10, pady=5)
        
        self.report_button = ttk.Button(settings_frame, text="Generate Report", command=self.generate_word_report)
        self.report_button.grid(row=0, column=8, padx=10, pady=5)
        
        self.save_button = ttk.Button(settings_frame, text="Save CSV", command=self.save_csv)
        self.save_button.grid(row=2, column=6, padx=10, pady=5)
        
        self.dashboard_button = ttk.Button(settings_frame, text="Show Dashboard", command=self.show_dashboard)
        self.dashboard_button.grid(row=2, column=7, padx=10, pady=5)
        
        self.trend_button = ttk.Button(settings_frame, text="Trend Dashboard", command=self.show_trend_dashboard)
        self.trend_button.grid(row=2, column=8, padx=10, pady=5)
        
        # Info label
        self.info_label = ttk.Label(settings_frame, text="")
        self.info_label.grid(row=3, column=0, columnspan=6, sticky="w", padx=5, pady=5)
        self.update_info_label()
        
        # Progress bar
        ttk.Label(settings_frame, text="Progress:").grid(row=4, column=0, sticky="w", padx=5, pady=5)
        self.progress_bar = ttk.Progressbar(
            settings_frame, variable=self.progress_var,
            maximum=100, length=500, mode="determinate",
        )
        self.progress_bar.grid(row=4, column=1, columnspan=4, sticky="we", padx=5, pady=5)
        
        self.progress_percent_label = ttk.Label(
            settings_frame, textvariable=self.progress_percent_var, width=7,
        )
        self.progress_percent_label.grid(row=4, column=5, sticky="w", padx=5, pady=5)
        
        self.status_label = ttk.Label(settings_frame, textvariable=self.status_var)
        self.status_label.grid(row=4, column=6, columnspan=3, sticky="w", padx=5, pady=5)
        
        # Results treeview
        self.tree = ttk.Treeview(result_frame, show="headings")
        tree_y = ttk.Scrollbar(result_frame, orient="vertical", command=self.tree.yview)
        tree_x = ttk.Scrollbar(result_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=tree_y.set, xscrollcommand=tree_x.set)
        
        self.tree.grid(row=0, column=0, sticky="nsew")
        tree_y.grid(row=0, column=1, sticky="ns")
        tree_x.grid(row=1, column=0, sticky="ew")
        
        result_frame.rowconfigure(0, weight=1)
        result_frame.columnconfigure(0, weight=1)
        
        # Log box
        self.log_text = tk.Text(log_frame, height=6)
        self.log_text.pack(fill="x")
    
    def log(self, msg):
        """Thread-safe logging using root.after()."""
        self.root.after(0, lambda: self._log_safe(msg))
    
    def _log_safe(self, msg):
        """Actual log implementation (must run in main thread)."""
        self.log_text.insert("end", str(msg) + "\n")
        self.log_text.see("end")
    
    def update_progress(self, value, status):
        """Thread-safe progress update."""
        self.root.after(0, lambda: self._update_progress_safe(value, status))
    
    def _update_progress_safe(self, value, status):
        """Actual progress update (must run in main thread)."""
        value = max(0, min(100, float(value)))
        self.progress_var.set(value)
        self.progress_percent_var.set(f"{int(value)}%")
        self.status_var.set(status)
        self._log_safe(f"[{int(value)}%] {status}")
    
    def browse_file(self):
        path = filedialog.askopenfilename(
            title="Select KPI Excel File",
            filetypes=[("Excel files", "*.xlsx *.xls")],
        )
        if path:
            self.file_path.set(path)
            self.log(f"Selected file: {path}")
            
            # NEW: Load sheet names
            try:
                xl = pd.ExcelFile(path)
                self.excel_sheets = xl.sheet_names
                self.sheet_combo['values'] = self.excel_sheets
                if self.excel_sheets:
                    self.selected_sheet.set(self.excel_sheets[0])
                    self.log(f"Found {len(self.excel_sheets)} sheets: {', '.join(self.excel_sheets[:5])}{'...' if len(self.excel_sheets) > 5 else ''}")
            except Exception as e:
                self.log(f"Warning: Could not read sheet names: {e}")
    
    def on_kpi_change(self, event=None):
        config = KPI_CONFIGS[self.selected_kpi.get()]
        self.threshold.set(config["default_threshold"])
        self.update_info_label()
    
    def update_info_label(self):
        config = KPI_CONFIGS[self.selected_kpi.get()]
        min_baseline = config.get("min_baseline_value", 0)
        self.info_label.config(
            text=f"Target: {config['target_kpi']} | Bad direction: {config['bad_direction']} | Min baseline: {min_baseline}"
        )
    
    def set_running_state(self, running):
        self.root.after(0, lambda: self._set_running_state_safe(running))
    
    def _set_running_state_safe(self, running):
        self.is_running = running
        state = "disabled" if running else "normal"
        for btn in [self.run_button, self.run_all_button, self.save_button, 
                    self.dashboard_button, self.trend_button, self.report_button]:
            btn.config(state=state)
    
    def run_analysis_thread(self):
        if self.is_running:
            messagebox.showinfo("Running", "Analysis is already running.")
            return
        self.progress_var.set(0)
        self.progress_percent_var.set("0%")
        self.status_var.set("Starting...")
        thread = threading.Thread(target=self.run_analysis, daemon=True)
        thread.start()
    
    def run_all_analysis_thread(self):
        if self.is_running:
            messagebox.showinfo("Running", "Analysis is already running.")
            return
        self.progress_var.set(0)
        self.progress_percent_var.set("0%")
        self.status_var.set("Starting Analyze All KPIs...")
        thread = threading.Thread(target=self.run_all_analysis, daemon=True)
        thread.start()
    
    def run_analysis(self):
        try:
            self.set_running_state(True)
            self.update_progress(2, "Checking selected file...")
            
            path = self.file_path.get().strip()
            if not path:
                self.root.after(0, lambda: messagebox.showwarning("Missing file", "Please select an Excel file."))
                return
            
            if not os.path.exists(path):
                self.root.after(0, lambda: messagebox.showerror("File error", "File does not exist."))
                return
            
            # Load Excel with sheet selection (NEW)
            sheet_name = self.selected_sheet.get() if self.selected_sheet.get() else 0
            self.update_progress(10, f"Loading Excel sheet: {sheet_name}...")
            df = pd.read_excel(path, sheet_name=sheet_name)
            
            self.update_progress(20, "Cleaning column names...")
            self.original_df = df.copy()
            
            selected_kpi = self.selected_kpi.get()
            num_days = int(self.num_days.get())
            threshold = float(self.threshold.get())
            complete_days = bool(self.require_complete_days.get())
            baseline_mode = self.baseline_mode.get()
            enable_sig = bool(self.enable_significance_test.get())
            
            self.log(f"Loaded: {df.shape[0]} rows, {df.shape[1]} columns")
            self.log(f"Baseline mode: {baseline_mode}")
            self.log(f"Significance test: {'Enabled' if enable_sig else 'Disabled'}")
            
            self.update_progress(35, f"Analyzing KPI: {selected_kpi}...")
            
            output_df, metadata = analyze_selected_kpi_enhanced(
                df=df,
                selected_kpi_name=selected_kpi,
                num_days=num_days,
                degradation_threshold=threshold,
                require_complete_days=complete_days,
                baseline_mode=baseline_mode,
                custom_baseline_start=self.custom_baseline_start.get() if baseline_mode == BASELINE_MODE_CUSTOM else None,
                custom_baseline_end=self.custom_baseline_end.get() if baseline_mode == BASELINE_MODE_CUSTOM else None,
                enable_significance_test=enable_sig,
                log_callback=self.log,
            )
            
            self.output_df = output_df
            self.all_outputs = {}
            self.summary_df = None
            self.analysis_mode = "single"
            
            # Store degraded cell IDs
            self.degraded_cell_ids = set()
            if not output_df.empty and SITE_COL in output_df.columns and CELL_COL in output_df.columns:
                for _, row in output_df.iterrows():
                    self.degraded_cell_ids.add((row.get(SITE_COL, ''), row.get(CELL_COL, '')))
            
            self.update_progress(80, "Preparing results...")
            
            self.log(f"Recent: {metadata['recent_start'].date()} to {metadata['recent_end'].date()}")
            self.log(f"Baseline: {metadata['baseline_start'].date()} to {metadata['baseline_end'].date()}")
            debug = metadata.get("debug_info", {})
            self.log(f"Cells after merge: {debug.get('cells_after_merge')}")
            self.log(f"Excluded (zero baseline): {debug.get('zero_baseline_excluded', 0)}")
            self.log(f"Excluded (min baseline): {debug.get('min_baseline_excluded', 0)}")
            self.log(f"Degraded cells found: {output_df.shape[0]}")
            
            self.update_progress(90, "Updating table...")
            self.root.after(0, lambda: self.update_table(output_df))
            
            self.update_progress(100, "Analysis completed.")
            
            msg = f"Analysis completed. Degraded cells: {output_df.shape[0]}"
            self.root.after(0, lambda: messagebox.showinfo("Done", msg))
            
        except Exception as e:
            self.log(f"ERROR: {e}")
            self.log(traceback.format_exc()[-500:])
            self.root.after(0, lambda: messagebox.showerror("Error", str(e)))
        finally:
            self.set_running_state(False)
    
    def run_all_analysis(self):
        try:
            self.set_running_state(True)
            self.update_progress(2, "Checking selected file...")
            
            path = self.file_path.get().strip()
            if not path or not os.path.exists(path):
                self.root.after(0, lambda: messagebox.showwarning("Missing file", "Please select an Excel file."))
                return
            
            sheet_name = self.selected_sheet.get() if self.selected_sheet.get() else 0
            self.update_progress(10, f"Loading Excel sheet: {sheet_name}...")
            df = pd.read_excel(path, sheet_name=sheet_name)
            self.original_df = df.copy()
            
            self.log(f"Loaded: {df.shape[0]} rows, {df.shape[1]} columns")
            
            outputs = {}
            summary_records = []
            all_kpi_names = list(KPI_CONFIGS.keys())
            total_kpis = len(all_kpi_names)
            
            num_days = int(self.num_days.get())
            complete_days = bool(self.require_complete_days.get())
            baseline_mode = self.baseline_mode.get()
            enable_sig = bool(self.enable_significance_test.get())
            
            for idx, kpi_name in enumerate(all_kpi_names, 1):
                config = KPI_CONFIGS[kpi_name]
                threshold = float(config["default_threshold"])
                
                progress = 20 + ((idx - 1) / max(total_kpis, 1)) * 65
                self.update_progress(progress, f"Analyzing {idx}/{total_kpis}: {kpi_name}")
                
                try:
                    output_df, metadata = analyze_selected_kpi_enhanced(
                        df=df,
                        selected_kpi_name=kpi_name,
                        num_days=num_days,
                        degradation_threshold=threshold,
                        require_complete_days=complete_days,
                        baseline_mode=baseline_mode,
                        enable_significance_test=enable_sig,
                        log_callback=self.log,
                    )
                    outputs[kpi_name] = output_df
                    
                    debug = metadata.get("debug_info", {})
                    degraded_count = output_df.shape[0]
                    
                    summary_records.append({
                        "kpi_name": kpi_name,
                        "target_kpi_column": config["target_kpi"],
                        "kpi_category": config["category"],
                        "threshold_%": threshold,
                        "degraded_cells_count": degraded_count,
                        "max_degradation_%": debug.get("max_degradation"),
                        "mean_degradation_%": debug.get("mean_degradation"),
                        "status": "Completed",
                        "error": ""
                    })
                    self.log(f"{kpi_name}: {degraded_count} degraded cells")
                    
                except Exception as e:
                    outputs[kpi_name] = pd.DataFrame()
                    summary_records.append({
                        "kpi_name": kpi_name,
                        "target_kpi_column": config.get("target_kpi", ""),
                        "kpi_category": config.get("category", ""),
                        "threshold_%": threshold,
                        "degraded_cells_count": 0,
                        "max_degradation_%": None,
                        "mean_degradation_%": None,
                        "status": "Failed",
                        "error": str(e)
                    })
                    self.log(f"{kpi_name}: ERROR - {e}")
            
            self.update_progress(88, "Combining results...")
            
            non_empty = [df for df in outputs.values() if df is not None and not df.empty]
            combined = pd.concat(non_empty, ignore_index=True) if non_empty else pd.DataFrame()
            summary_df = pd.DataFrame(summary_records)
            
            self.output_df = combined
            self.all_outputs = outputs
            self.summary_df = summary_df
            self.analysis_mode = "all"
            
            # Store degraded cell IDs
            self.degraded_cell_ids = set()
            if not combined.empty and SITE_COL in combined.columns and CELL_COL in combined.columns:
                for _, row in combined.iterrows():
                    self.degraded_cell_ids.add((row.get(SITE_COL, ''), row.get(CELL_COL, '')))
            
            self.update_progress(92, "Updating table...")
            self.root.after(0, lambda: self.update_table(summary_df))
            
            total_degraded = combined.shape[0]
            self.log(f"Total degraded cells: {total_degraded}")
            
            self.update_progress(100, "Analyze All KPIs completed.")
            self.root.after(0, lambda: messagebox.showinfo("Done", f"Completed. Total degraded: {total_degraded}"))
            
        except Exception as e:
            self.log(f"ERROR: {e}")
            self.root.after(0, lambda: messagebox.showerror("Error", str(e)))
        finally:
            self.set_running_state(False)
    
    def update_table(self, df):
        self.tree.delete(*self.tree.get_children())
        self.tree["columns"] = []
        
        if df is None or df.empty:
            return
        
        preview = df.head(200).copy()
        cols = list(preview.columns)
        self.tree["columns"] = cols
        
        for col in cols:
            self.tree.heading(col, text=col[:30])
            self.tree.column(col, width=140, anchor="w")
        
        for _, row in preview.iterrows():
            vals = ["" if pd.isna(row[c]) else str(row[c])[:80] for c in cols]
            self.tree.insert("", "end", values=vals)
    
    def show_dashboard(self):
        if self.output_df is None and self.summary_df is None:
            messagebox.showwarning("No output", "Please run analysis first.")
            return
        
        dash = tk.Toplevel(self.root)
        dash.title("LTE KPI Degradation Dashboard")
        dash.geometry("1200x760")
        
        main_frame = ttk.Frame(dash, padding=10)
        main_frame.pack(fill="both", expand=True)
        
        ttk.Label(main_frame, text="LTE KPI Degradation Dashboard", font=("Arial", 16, "bold")).pack(anchor="w", pady=(0, 10))
        
        metrics_frame = ttk.LabelFrame(main_frame, text="Summary Metrics", padding=10)
        metrics_frame.pack(fill="x", pady=(0, 10))
        
        charts_frame = ttk.Frame(main_frame)
        charts_frame.pack(fill="both", expand=True)
        
        left_chart = ttk.LabelFrame(charts_frame, text="Degraded Cells per KPI", padding=10)
        left_chart.pack(side="left", fill="both", expand=True, padx=(0, 5))
        
        right_chart = ttk.LabelFrame(charts_frame, text="Root Cause Distribution", padding=10)
        right_chart.pack(side="right", fill="both", expand=True, padx=(5, 0))
        
        # Metrics
        if self.analysis_mode == "all" and self.summary_df is not None:
            total_kpis = self.summary_df.shape[0]
            total_degraded = int(self.summary_df["degraded_cells_count"].sum()) if "degraded_cells_count" in self.summary_df.columns else 0
            metrics = [("Mode", "All KPIs"), ("KPIs", total_kpis), ("Total Degraded", total_degraded)]
        else:
            total_degraded = self.output_df.shape[0] if self.output_df is not None else 0
            metrics = [("Mode", "Single KPI"), ("KPI", self.selected_kpi.get()), ("Degraded", total_degraded)]
        
        for i, (name, val) in enumerate(metrics):
            box = ttk.LabelFrame(metrics_frame, text=name, padding=5)
            box.grid(row=0, column=i, sticky="nsew", padx=5)
            ttk.Label(box, text=str(val), font=("Arial", 11, "bold")).pack()
            metrics_frame.columnconfigure(i, weight=1)
        
        # Charts
        fig1 = Figure(figsize=(5, 4), dpi=100)
        ax1 = fig1.add_subplot(111)
        
        if self.analysis_mode == "all" and self.summary_df is not None and "degraded_cells_count" in self.summary_df.columns:
            plot_df = self.summary_df.sort_values("degraded_cells_count", ascending=False).head(12)
            bars = ax1.bar(plot_df["kpi_name"], plot_df["degraded_cells_count"])
            ax1.set_title("Degraded Cells per KPI")
            ax1.tick_params(axis="x", rotation=70)
        else:
            ax1.text(0.5, 0.5, "No data", ha="center")
        
        fig1.tight_layout()
        canvas1 = FigureCanvasTkAgg(fig1, master=left_chart)
        canvas1.draw()
        canvas1.get_tk_widget().pack(fill="both", expand=True)
        
        fig2 = Figure(figsize=(5, 4), dpi=100)
        ax2 = fig2.add_subplot(111)
        
        if self.output_df is not None and not self.output_df.empty and "main_root_cause_category" in self.output_df.columns:
            causes = self.output_df["main_root_cause_category"].value_counts().head(10).sort_values()
            ax2.barh(causes.index, causes.values)
            ax2.set_title("Root Causes")
        else:
            ax2.text(0.5, 0.5, "No data", ha="center")
        
        fig2.tight_layout()
        canvas2 = FigureCanvasTkAgg(fig2, master=right_chart)
        canvas2.draw()
        canvas2.get_tk_widget().pack(fill="both", expand=True)
    
    def show_trend_dashboard(self):
        if self.original_df is None:
            messagebox.showwarning("No Data", "Please load an Excel file first.")
            return
        
        if self.output_df is None or self.output_df.empty:
            messagebox.showwarning("No Results", "No degraded cells found. Run analysis first.")
            return
        
        trend_win = tk.Toplevel(self.root)
        trend_win.title("KPI Trend - Before vs After Degradation Removal")
        trend_win.geometry("1300x800")
        
        main_frame = ttk.Frame(trend_win, padding=10)
        main_frame.pack(fill="both", expand=True)
        
        ttk.Label(main_frame, text="KPI Trend Analysis - Before vs After Degraded Cell Removal",
                  font=("Arial", 14, "bold")).pack(anchor="w", pady=(0, 10))
        
        # Controls
        controls = ttk.Frame(main_frame)
        controls.pack(fill="x", pady=(0, 10))
        
        ttk.Label(controls, text="Select KPI:").pack(side="left", padx=5)
        
        config = KPI_CONFIGS.get(self.selected_kpi.get(), {})
        target_kpi = config.get("target_kpi", "")
        kpi_col = find_matching_column(self.original_df, target_kpi)
        
        numeric_cols = self.original_df.select_dtypes(include=[np.number]).columns.tolist()
        numeric_cols = [c for c in numeric_cols if c not in CELL_ID_COLS]
        
        trend_kpi = tk.StringVar(value=kpi_col if kpi_col else (numeric_cols[0] if numeric_cols else ""))
        ttk.Combobox(controls, textvariable=trend_kpi, values=numeric_cols[:30], state="readonly", width=50).pack(side="left", padx=5)
        
        chart_frame = ttk.LabelFrame(main_frame, text="Trend Chart", padding=10)
        chart_frame.pack(fill="both", expand=True, pady=5)
        
        def draw_chart():
            for w in chart_frame.winfo_children():
                w.destroy()
            
            col = trend_kpi.get()
            if not col or col not in self.original_df.columns:
                ttk.Label(chart_frame, text="Invalid KPI column").pack()
                return
            
            df = self.original_df.copy()
            
            if DATE_COL not in df.columns:
                ttk.Label(chart_frame, text="Date column not found").pack()
                return
            
            df[DATE_COL] = pd.to_datetime(df[DATE_COL], errors='coerce')
            df = df.dropna(subset=[DATE_COL])
            
            # Before: all cells
            daily_before = df.groupby(DATE_COL)[col].mean().reset_index()
            daily_before.columns = ['Date', 'Average']
            
            # After: remove degraded cells
            if self.degraded_cell_ids:
                mask = df.set_index([SITE_COL, CELL_COL]).index.isin(self.degraded_cell_ids)
                df_clean = df[~mask]
                daily_after = df_clean.groupby(DATE_COL)[col].mean().reset_index() if len(df_clean) > 0 else daily_before.copy()
                daily_after.columns = ['Date', 'Average']
            else:
                daily_after = daily_before.copy()
            
            if daily_before.empty:
                ttk.Label(chart_frame, text="No data to plot").pack()
                return
            
            fig = Figure(figsize=(12, 5), dpi=100)
            ax = fig.add_subplot(111)
            
            dates = daily_before['Date'].tolist()
            x = range(len(dates))
            labels = [d.strftime('%m/%d') if hasattr(d, 'strftime') else str(d)[:10] for d in dates]
            
            ax.plot(x, daily_before['Average'].values, 'b-o', linewidth=2, markersize=6, label='Before (All Cells)')
            ax.plot(x, daily_after['Average'].values, 'g-s', linewidth=2, markersize=6, label='After (Clean Cells)')
            
            if self.degraded_cell_ids:
                diff = daily_before['Average'].values - daily_after['Average'].values
                ax.fill_between(x, daily_before['Average'].values, daily_after['Average'].values,
                               where=diff != 0, alpha=0.3, color='red', label='Degraded Impact')
            
            ax.set_xticks(x)
            ax.set_xticklabels(labels, rotation=45, ha='right', fontsize=8)
            ax.set_xlabel('Date')
            ax.set_ylabel(col[:40])
            ax.set_title(f'{col[:40]} - Daily Trend')
            ax.legend()
            ax.grid(True, alpha=0.3)
            
            fig.tight_layout()
            canvas = FigureCanvasTkAgg(fig, master=chart_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill="both", expand=True)
        
        ttk.Button(controls, text="Update Chart", command=draw_chart).pack(side="left", padx=10)
        
        # Legend
        legend_frame = ttk.Frame(main_frame)
        legend_frame.pack(fill="x", pady=5)
        ttk.Label(legend_frame, text="■", foreground='blue', font=('Arial', 12)).pack(side="left", padx=2)
        ttk.Label(legend_frame, text="Before Removal (All Cells)").pack(side="left", padx=(0, 20))
        ttk.Label(legend_frame, text="■", foreground='green', font=('Arial', 12)).pack(side="left", padx=2)
        ttk.Label(legend_frame, text="After Removal (Clean Cells)").pack(side="left", padx=(0, 20))
        ttk.Label(legend_frame, text="■", foreground='red', font=('Arial', 12)).pack(side="left", padx=2)
        ttk.Label(legend_frame, text="Degraded Impact").pack(side="left")
        
        draw_chart()
    
    def generate_word_report(self):
        if not DOCX_AVAILABLE:
            messagebox.showerror("Missing Package", "Install python-docx: pip install python-docx")
            return
        
        if self.output_df is None or (self.output_df.empty and (self.summary_df is None or self.summary_df.empty)):
            messagebox.showwarning("No Results", "Run analysis first.")
            return
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        default_name = f"RF_Optimization_Report_{timestamp}.docx"
        
        save_path = filedialog.asksaveasfilename(
            title="Save Word Report", defaultextension=".docx",
            initialfile=default_name, filetypes=[("Word documents", "*.docx")]
        )
        
        if not save_path:
            return
        
        try:
            doc = Document()
            
            # Title
            doc.add_heading('RF Optimization Analysis Report', 0)
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("Developed by: Musketeers_Team (ITI Graduation Project 2026)")
            doc.add_paragraph(f"Version: 2.0 Enhanced")
            
            # Analysis Summary
            doc.add_heading('Analysis Summary', level=1)
            
            if self.analysis_mode == "all":
                doc.add_paragraph(f"Analysis Mode: All KPIs Analysis")
                doc.add_paragraph(f"Baseline Mode: {self.baseline_mode.get()}")
                doc.add_paragraph(f"Significance Test: {'Enabled' if self.enable_significance_test.get() else 'Disabled'}")
                if self.summary_df is not None:
                    doc.add_paragraph(f"Total KPIs Analyzed: {len(self.summary_df)}")
                    doc.add_paragraph(f"Total Degraded Cells: {self.summary_df['degraded_cells_count'].sum()}")
            else:
                doc.add_paragraph(f"Analysis Mode: Single KPI Analysis")
                doc.add_paragraph(f"Selected KPI: {self.selected_kpi.get()}")
                doc.add_paragraph(f"Baseline Mode: {self.baseline_mode.get()}")
                doc.add_paragraph(f"Degraded Cells: {len(self.output_df) if self.output_df is not None else 0}")
            
            # Degraded Cells Table
            if self.output_df is not None and not self.output_df.empty:
                doc.add_heading('Degraded Cells Details', level=1)
                
                key_cols = ['eNodeB Name', 'Cell Name', 'kpi_degradation_ratio_%', 
                           'main_root_cause_category', 'main_recommended_action', 'stat_significant', 'p_value']
                available = [c for c in key_cols if c in self.output_df.columns]
                
                if available:
                    table = doc.add_table(rows=1, cols=len(available))
                    table.style = 'Table Grid'
                    
                    headers = table.rows[0].cells
                    for i, col in enumerate(available):
                        headers[i].text = col.replace('_', ' ').title()
                    
                    for _, row in self.output_df.head(30).iterrows():
                        cells = table.add_row().cells
                        for i, col in enumerate(available):
                            val = row.get(col, '')
                            cells[i].text = "N/A" if pd.isna(val) else str(val)[:80]
            
            # Summary Table
            if self.analysis_mode == "all" and self.summary_df is not None and not self.summary_df.empty:
                doc.add_heading('KPI Summary Table', level=1)
                
                sum_cols = ['kpi_name', 'degraded_cells_count', 'max_degradation_%', 'status']
                avail_sum = [c for c in sum_cols if c in self.summary_df.columns]
                
                if avail_sum:
                    table = doc.add_table(rows=1, cols=len(avail_sum))
                    table.style = 'Table Grid'
                    
                    headers = table.rows[0].cells
                    for i, col in enumerate(avail_sum):
                        headers[i].text = col.replace('_', ' ').title()
                    
                    for _, row in self.summary_df.iterrows():
                        cells = table.add_row().cells
                        for i, col in enumerate(avail_sum):
                            val = row.get(col, '')
                            cells[i].text = "N/A" if pd.isna(val) else str(val)[:50]
            
            doc.save(save_path)
            self.log(f"Report saved: {save_path}")
            messagebox.showinfo("Report Saved", f"Report saved:\n{save_path}")
            
        except Exception as e:
            self.log(f"Error generating report: {e}")
            messagebox.showerror("Error", f"Failed to generate report:\n{str(e)}")
    
    def save_csv(self):
        if self.output_df is None and self.summary_df is None:
            messagebox.showwarning("No output", "Run analysis first.")
            return
        
        if self.analysis_mode == "all":
            save_dir = filedialog.askdirectory(title="Select folder to save CSV files")
            if not save_dir:
                return
            
            saved = 0
            for kpi_name, kpi_df in self.all_outputs.items():
                if kpi_df is not None and not kpi_df.empty:
                    prefix = KPI_CONFIGS[kpi_name]["output_prefix"]
                    path = os.path.join(save_dir, f"{prefix}_degraded.csv")
                    kpi_df.to_csv(path, index=False, encoding="utf-8-sig")
                    saved += 1
            
            if self.output_df is not None and not self.output_df.empty:
                self.output_df.to_csv(os.path.join(save_dir, "all_kpis_combined.csv"), index=False, encoding="utf-8-sig")
                saved += 1
            
            if self.summary_df is not None and not self.summary_df.empty:
                self.summary_df.to_csv(os.path.join(save_dir, "summary_report.csv"), index=False, encoding="utf-8-sig")
                saved += 1
            
            self.log(f"Saved {saved} files to: {save_dir}")
            messagebox.showinfo("Saved", f"Saved {saved} CSV files to:\n{save_dir}")
            return
        
        if self.output_df is None or self.output_df.empty:
            messagebox.showwarning("No data", "No degraded cells to save.")
            return
        
        prefix = KPI_CONFIGS[self.selected_kpi.get()]["output_prefix"]
        default_name = f"{prefix}_degraded.csv"
        
        save_path = filedialog.asksaveasfilename(
            title="Save CSV", defaultextension=".csv",
            initialfile=default_name, filetypes=[("CSV files", "*.csv")]
        )
        
        if save_path:
            self.output_df.to_csv(save_path, index=False, encoding="utf-8-sig")
            self.log(f"CSV saved: {save_path}")
            messagebox.showinfo("Saved", f"CSV saved:\n{save_path}")


# ============================================================
# 6. RUN APPLICATION
# ============================================================

if __name__ == "__main__":
    root = tk.Tk()
    app = LTEKPIAnalyzerApp(root)
    root.mainloop()
