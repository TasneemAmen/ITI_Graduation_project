# ============================================================
# LTE KPI Degradation Analyzer - Generate Word Report
# ============================================================
# This file contains functions for generating Word reports.
# ============================================================

import pandas as pd
from datetime import datetime

# Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def generate_word_report(output_df, summary_df, analysis_mode, selected_kpi, baseline_mode, 
                          enable_significance_test, save_path, log_callback=None):
    """
    Generate a Word report with analysis results.
    
    Args:
        output_df: Output DataFrame with degraded cells
        summary_df: Summary DataFrame (for all KPIs mode)
        analysis_mode: "single" or "all"
        selected_kpi: Currently selected KPI name
        baseline_mode: Baseline mode setting
        enable_significance_test: Whether t-test was enabled
        save_path: Path to save the report
        log_callback: Optional logging callback
        
    Returns:
        True if successful, False otherwise
    """
    def log_msg(msg):
        if log_callback:
            log_callback(msg)
    
    if not DOCX_AVAILABLE:
        log_msg("ERROR: python-docx not installed. Run: pip install python-docx")
        return False
    
    if output_df is None or (output_df.empty and (summary_df is None or summary_df.empty)):
        log_msg("ERROR: No results to generate report")
        return False
    
    try:
        doc = Document()
        
        # Title
        doc.add_heading('RF Optimization Analysis Report', 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("Developed by: Musketeers_Team (ITI Graduation Project 2026)")
        doc.add_paragraph(f"Version: 2.0 Enhanced")
        
        # Analysis Summary
        doc.add_heading('Analysis Summary', level=1)
        
        if analysis_mode == "all":
            doc.add_paragraph(f"Analysis Mode: All KPIs Analysis")
            doc.add_paragraph(f"Baseline Mode: {baseline_mode}")
            doc.add_paragraph(f"Significance Test: {'Enabled' if enable_significance_test else 'Disabled'}")
            if summary_df is not None:
                doc.add_paragraph(f"Total KPIs Analyzed: {len(summary_df)}")
                doc.add_paragraph(f"Total Degraded Cells: {summary_df['degraded_cells_count'].sum()}")
        else:
            doc.add_paragraph(f"Analysis Mode: Single KPI Analysis")
            doc.add_paragraph(f"Selected KPI: {selected_kpi}")
            doc.add_paragraph(f"Baseline Mode: {baseline_mode}")
            doc.add_paragraph(f"Degraded Cells: {len(output_df) if output_df is not None else 0}")
        
        # Degraded Cells Table
        if output_df is not None and not output_df.empty:
            doc.add_heading('Degraded Cells Details', level=1)
            
            key_cols = ['eNodeB Name', 'Cell Name', 'kpi_degradation_ratio_%', 
                       'main_root_cause_category', 'main_recommended_action', 'stat_significant', 'p_value']
            available = [c for c in key_cols if c in output_df.columns]
            
            if available:
                table = doc.add_table(rows=1, cols=len(available))
                table.style = 'Table Grid'
                
                headers = table.rows[0].cells
                for i, col in enumerate(available):
                    headers[i].text = col.replace('_', ' ').title()
                
                for _, row in output_df.head(30).iterrows():
                    cells = table.add_row().cells
                    for i, col in enumerate(available):
                        val = row.get(col, '')
                        cells[i].text = "N/A" if pd.isna(val) else str(val)[:80]
        
        # Summary Table
        if analysis_mode == "all" and summary_df is not None and not summary_df.empty:
            doc.add_heading('KPI Summary Table', level=1)
            
            sum_cols = ['kpi_name', 'degraded_cells_count', 'max_degradation_%', 'status']
            avail_sum = [c for c in sum_cols if c in summary_df.columns]
            
            if avail_sum:
                table = doc.add_table(rows=1, cols=len(avail_sum))
                table.style = 'Table Grid'
                
                headers = table.rows[0].cells
                for i, col in enumerate(avail_sum):
                    headers[i].text = col.replace('_', ' ').title()
                
                for _, row in summary_df.iterrows():
                    cells = table.add_row().cells
                    for i, col in enumerate(avail_sum):
                        val = row.get(col, '')
                        cells[i].text = "N/A" if pd.isna(val) else str(val)[:50]
        
        doc.save(save_path)
        log_msg(f"Report saved: {save_path}")
        return True
        
    except Exception as e:
        log_msg(f"Error generating report: {e}")
        return False
